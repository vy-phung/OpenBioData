import re
import os
import subprocess
import zipfile
from Bio import Entrez
try:
    from bs4 import BeautifulSoup as _BS4
except ImportError:
    _BS4 = None
try:
    from docx import Document
    import fitz
    import spacy
    from spacy.cli import download
except ImportError:
    Document = fitz = spacy = download = None
try:
    from NER.PDF import pdf
except Exception:
    pdf = None
try:
    from NER.WordDoc import wordDoc
except Exception:
    wordDoc = None
try:
    from NER.html import extractHTML
except Exception:
    extractHTML = None
try:
    from NER.word2Vec import word2vec
except Exception:
    word2vec = None
import urllib.parse, requests
from urllib.parse import urlparse as _urlparse, parse_qs as _parse_qs
from pathlib import Path

_DOWNLOAD_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
    "Referer": "https://www.google.com/",
}

def _github_blob_to_raw(url: str) -> str:
    """Convert a GitHub blob viewer URL to its raw content URL."""
    if "github.com" in url and "/blob/" in url:
        url = url.replace("github.com", "raw.githubusercontent.com")
        url = url.replace("/blob/", "/")
    return url

def _is_local_path(path: str) -> bool:
    """Return True if path is a local filesystem path (not a Drive folder ID)."""
    return bool(path and (path.startswith("/") or path.startswith("\\") or (len(path) > 1 and path[1] == ":")))

def _get_filename_from_url(url):
    """Extract the true filename from a URL, including query-param style like ?file=foo.docx."""
    parsed = _urlparse(url)
    qs = _parse_qs(parsed.query)
    if 'file' in qs:
        return qs['file'][0]
    name = os.path.basename(parsed.path)
    return name if name else 'downloaded_file'

def classify_url(url: str) -> str:
    """Return 'pdf', 'xlsx', 'docx', 'pptx', 'ppt', 'zip', or 'html' for a URL or local path.
    Checks query-param ?file= as well as the URL path extension."""
    parsed = _urlparse(url)
    qs = _parse_qs(parsed.query)
    file_param = qs.get("file", [""])[0].lower()
    from urllib.parse import unquote
    path_lower = unquote(parsed.path).lower()
    for candidate in [file_param, path_lower]:
        if not candidate:
            continue
        if candidate.endswith(".pdf"):
            return "pdf"
        if candidate.endswith((".xlsx", ".xls")):
            return "xlsx"
        if candidate.endswith(".docx") or candidate.endswith(".doc"):
            return "docx"
        if candidate.endswith(".pptx"):
            return "pptx"
        if candidate.endswith(".ppt"):
            return "ppt"
        if candidate.endswith(".zip"):
            return "zip"
    return "html"


def _extract_zip_text(zip_path, saveFolder=None) -> str:
    """Extract text from all readable files inside a ZIP archive.

    Handles nested PDFs, XLSX, DOCX, CSV/TSV, and plain-text files.
    """
    import tempfile as _tmpmod
    from pathlib import Path as _P
    parts = []
    try:
        with zipfile.ZipFile(str(zip_path)) as z:
            for entry_name in z.namelist():
                if entry_name.endswith("/"):
                    continue  # skip directories
                ext = _P(entry_name).suffix.lower()
                try:
                    data = z.read(entry_name)
                    tmp_dir = _tmpmod.gettempdir()
                    tmp_path = _P(tmp_dir) / _P(entry_name).name
                    tmp_path.write_bytes(data)
                    text = ""
                    if ext == ".pdf" and pdf is not None:
                        text = pdf.PDFFast(str(tmp_path), tmp_dir).extract_text()
                    elif ext in (".xlsx", ".xls"):
                        text = _extract_excel_text(tmp_path)
                    elif ext in (".docx", ".doc") and wordDoc is not None:
                        text = wordDoc.WordDocFast(str(tmp_path), tmp_dir).extractText()
                    elif ext in (".txt", ".csv", ".tsv"):
                        text = data.decode("utf-8", errors="replace")
                    elif ext in (".pptx", ".ppt"):
                        text = _extract_pptx(tmp_path)
                    if text:
                        parts.append(f"[ZIP entry: {entry_name}]\n{text}")
                except Exception as _ze:
                    print(f"❌ ZIP entry '{entry_name}' failed: {_ze}")
    except Exception as e:
        print(f"❌ ZIP extraction failed for {zip_path}: {e}")
    return "\n\n".join(parts)


def _extract_excel_text(xlsx_path) -> str:
    """Extract all cell text from an Excel file as a formatted string."""
    try:
        import pandas as _pd
        xls = _pd.ExcelFile(str(xlsx_path))
        parts = []
        for sheet in xls.sheet_names:
            df = _pd.read_excel(xls, sheet_name=sheet, header=None)
            df = df.fillna("").astype(str)
            rows = []
            for _, row in df.iterrows():
                cleaned = [c.strip() for c in row if c.strip()]
                if cleaned:
                    rows.append(" | ".join(cleaned))
            if rows:
                parts.append(f"[Sheet: {sheet}]\n" + "\n".join(rows))
        return "\n\n".join(parts) if parts else ""
    except Exception as e:
        print(f"❌ Excel text extraction failed: {e}")
        return ""


def _extract_pptx(path) -> str:
    """Extract text from a .pptx file (or .ppt converted to .pptx) using zipfile + XML parsing."""
    try:
        slides = []
        with zipfile.ZipFile(str(path)) as z:
            slide_files = sorted(
                n for n in z.namelist()
                if n.startswith("ppt/slides/slide") and n.endswith(".xml")
            )
            for i, sf in enumerate(slide_files, 1):
                content = z.read(sf).decode("utf-8", errors="replace")
                if _BS4:
                    soup = _BS4(content, "xml")
                    texts = [t.get_text(strip=True) for t in soup.find_all("a:t") if t.get_text(strip=True)]
                else:
                    texts = re.findall(r'<a:t[^>]*>([^<]+)</a:t>', content)
                if texts:
                    slides.append(f"[Slide {i}]\n" + " ".join(texts))
        return "\n\n".join(slides) if slides else ""
    except Exception as e:
        print(f"❌ PPTX extraction failed: {e}")
        return ""


def _convert_ppt_to_pptx(ppt_path) -> str:
    """Convert .ppt to .pptx via LibreOffice. Returns converted path or None."""
    from pathlib import Path as _Path
    ppt_path = _Path(ppt_path)
    pptx_path = ppt_path.with_suffix(".pptx")
    if pptx_path.exists():
        return str(pptx_path)
    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pptx",
             "--outdir", str(ppt_path.parent), str(ppt_path)],
            capture_output=True, timeout=60
        )
        if result.returncode == 0 and pptx_path.exists():
            return str(pptx_path)
    except Exception as e:
        print(f"⚠️ LibreOffice PPT conversion failed: {e}")
    return None


def _extract_ppt_text(ppt_path) -> str:
    """Extract text from a PPT or PPTX file."""
    from pathlib import Path as _Path
    ext = _Path(str(ppt_path)).suffix.lower()
    if ext == ".pptx":
        return _extract_pptx(ppt_path)
    converted = _convert_ppt_to_pptx(ppt_path)
    if converted:
        return _extract_pptx(converted)
    return _extract_pptx(ppt_path)


def _extract_docx_text(docx_path) -> str:
    """Extract text from a local .docx file using python-docx."""
    try:
        if Document is None:
            raise ImportError("python-docx not available")
        doc = Document(str(docx_path))
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        print(f"❌ DOCX text extraction failed: {e}")
        return ""


def extract_url_text(url: str, data_dir) -> dict:
    """Extract text from a single URL. Binary files are saved to data_dir.

    Returns:
        {"url": str, "name": str, "kind": str,
         "status": "ok"|"failed", "text": str, "error": str}
    """
    from pathlib import Path as _Path
    data_dir = _Path(data_dir)
    url = url.strip()
    name = _get_filename_from_url(url)
    kind = classify_url(url)
    result = {"url": url, "name": name, "kind": kind, "status": "ok", "text": "", "error": ""}

    if kind == "html":
        try:
            if extractHTML is not None:
                html = extractHTML.HTML(htmlContent=None, htmlLink=url, htmlFile="")
                text = html.getListSection()
            else:
                import requests as _req
                from bs4 import BeautifulSoup as _BS
                r = _req.get(url, headers=_DOWNLOAD_HEADERS, timeout=30)
                r.raise_for_status()
                soup = _BS(r.content, "html.parser")
                for tag in soup(["script", "style", "nav", "footer", "header"]):
                    tag.decompose()
                text = soup.get_text(separator="\n", strip=True)
            if text:
                result["text"] = text
            else:
                result["status"] = "failed"
                result["error"] = "No readable text found"
        except Exception as e:
            result["status"] = "failed"
            result["error"] = str(e)
    else:
        dest = data_dir / name
        try:
            if not dest.exists():
                content = _download_file(url)
                with open(dest, "wb") as f:
                    f.write(content)
            # Extract text from the downloaded file
            if kind == "pdf" and pdf is not None:
                text = pdf.PDFFast(str(dest), str(data_dir)).extract_text()
            elif kind == "docx":
                if wordDoc is not None:
                    text = wordDoc.WordDocFast(str(dest), str(data_dir)).extractText()
                else:
                    text = _extract_docx_text(dest)
            elif kind == "xlsx":
                text = _extract_excel_text(dest)
            elif kind in ("ppt", "pptx"):
                text = _extract_ppt_text(dest)
            else:
                text = ""

            if text:
                result["text"] = text
            else:
                result["status"] = "failed"
                result["error"] = f"No text extracted from {name}"
        except Exception as e:
            result["status"] = "failed"
            result["error"] = str(e)

    return result


def process_sources_to_docx(links: list, output_path, dataset_label: str = "") -> tuple:
    """Extract text from each URL, write a formatted DOCX to output_path.

    Returns (results: list[dict], failed_links: list[(url, reason)])
    """
    from pathlib import Path as _Path
    output_path = _Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    data_dir = output_path.parent

    if Document is None:
        print("⚠️ python-docx not available — cannot write DOCX")
        return [], []

    doc = Document()
    doc.add_heading("Supplementary Material Extract", level=0)
    if dataset_label:
        doc.add_paragraph(f"Source dataset: {dataset_label}")
    doc.add_paragraph(f"Total links processed: {len(links)}")

    results = []
    failed_links = []

    for idx, url in enumerate(links, 1):
        url = url.strip()
        if not url:
            continue
        r = extract_url_text(url, data_dir)
        results.append(r)

        p = doc.add_paragraph()
        run = p.add_run(f"The source - {r['name']}")
        run.bold = True
        doc.add_paragraph(f"URL: {url}")
        doc.add_paragraph(f"Type: {r['kind'].upper()}")

        if r["status"] == "ok" and r["text"]:
            for line in r["text"].split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
        else:
            err = r.get("error") or "Unknown error"
            doc.add_paragraph(f"[Could not extract content: {err}]")
            failed_links.append((url, err))

        doc.add_paragraph("─" * 80)

    doc.save(str(output_path))
    print(f"✅ Document saved: {output_path}")
    return results, failed_links


def _download_file(url):
    """Download a file with browser-like headers. Raises ValueError if response is HTML (blocked/403)."""
    r = requests.get(url, headers=_DOWNLOAD_HEADERS, timeout=30, allow_redirects=True)
    r.raise_for_status()
    ct = r.headers.get('Content-Type', '')
    if 'text/html' in ct:
        raise ValueError(
            f"Publisher returned an HTML page instead of the file — it likely requires a subscription "
            f"or blocks server/cloud IP addresses. URL: {url}"
        )
    return r.content
import pandas as pd
try:
    import model
except Exception:
    model = None
try:
    import pipeline
except Exception:
    pipeline = None
import tempfile
import nltk
nltk.download('punkt_tab')
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    try: 
        nltk.download('stopwords')
    except:
        print("have to use our own created stopword")
        STOPWORDS = {
            "the","a","an","in","on","of","and","or","for","with","to","from",
            "is","are","was","were","be","been","by","this","that","these","those",
            "it","its","as","at","but","not","no","so","if","their","there","about",
            "into","such","than","other","then","also","can","may","might","should"
        }

def download_excel_file(url, save_path="temp.xlsx"):
    if "view.officeapps.live.com" in url:
        parsed_url = urllib.parse.parse_qs(urllib.parse.urlparse(url).query)
        real_url = urllib.parse.unquote(parsed_url["src"][0])
        content = _download_file(real_url)
        with open(save_path, "wb") as f:
            f.write(content)
        return save_path
    elif url.startswith("http"):
        filename = _get_filename_from_url(url)
        if filename.lower().endswith(('.xls', '.xlsx')):
            try:
                content = _download_file(url)
                with open(save_path, "wb") as f:
                    f.write(content)
                    print(len(content))
                return save_path
            except Exception as e:
                print(f"❌ Excel download failed: {e}")
                return url
    print("URL must point directly to an .xls or .xlsx file or it is already downloaded.")
    return url
        
from pathlib import Path
import pandas as pd

def process_file(link, saveFolder):
    """Returns (file_type, full_path, name) for a given link."""
    name = Path(link).name
    ext = Path(name).suffix.lower()
    file_path = Path(saveFolder) / name

    # If it's already in saveFolder, update link to local path
    if file_path.is_file():
        link = str(file_path)

    return ext, link, file_path

import asyncio
import aiohttp
_html_cache = {}

async def async_fetch_html(link: str, timeout: int = 15) -> str:
    """Fetch HTML asynchronously with caching."""
    if link in _html_cache:
        return _html_cache[link]

    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(link, timeout=timeout) as resp:
                if resp.status != 200:
                    print(f"⚠️ Failed {link} ({resp.status})")
                    return ""
                html_content = await resp.text()
                _html_cache[link] = html_content
                return html_content
    except Exception as e:
        print(f"❌ async_fetch_html error for {link}: {e}")
        return ""

async def ensure_local_file(link: str, saveFolder: str) -> str:
    """Ensure file is available locally (Drive or web). Returns local path."""
    link = _github_blob_to_raw(link)
    name = _get_filename_from_url(link)
    local_temp_path = os.path.join(tempfile.gettempdir(), name)

    if os.path.exists(local_temp_path):
        return local_temp_path

    # Try Drive only when saveFolder is a real Drive folder ID (not a local temp path)
    _use_drive = pipeline is not None and not _is_local_path(saveFolder)
    file_id = await asyncio.to_thread(pipeline.find_drive_file, name, saveFolder) if _use_drive else None
    if file_id:
        await asyncio.to_thread(pipeline.download_file_from_drive, name, saveFolder, local_temp_path)
    else:
        # Web download asynchronously with browser headers
        async with aiohttp.ClientSession(headers=_DOWNLOAD_HEADERS) as session:
            async with session.get(link, timeout=20, allow_redirects=True) as resp:
                resp.raise_for_status()
                ct = resp.headers.get('Content-Type', '')
                if 'text/html' in ct:
                    raise ValueError(
                        f"Publisher returned HTML instead of file (blocked or requires login): {link}"
                    )
                content = await resp.read()
                with open(local_temp_path, "wb") as f:
                    f.write(content)
        # Upload back to Drive only when we have a real Drive folder ID
        if _use_drive:
            await asyncio.to_thread(pipeline.upload_file_to_drive, local_temp_path, name, saveFolder)

    return local_temp_path

async def async_extract_text(link, saveFolder):
    """Extract text from any URL or local path.

    Detects type via classify_url() — handles pdf, docx, xlsx, ppt/pptx, and html.
    Binary files are downloaded to saveFolder before extraction.
    """
    try:
        link = _github_blob_to_raw(link)
        # Non-URLs (local paths) — check file extension directly
        if not link.startswith("http"):
            kind = classify_url(link)
        else:
            kind = classify_url(link)

        if kind == "pdf":
            if pdf is None:
                return ""
            local_path = await ensure_local_file(link, saveFolder)
            return await asyncio.to_thread(lambda: pdf.PDFFast(local_path, saveFolder).extract_text())

        elif kind == "docx":
            if wordDoc is None:
                return ""
            local_path = await ensure_local_file(link, saveFolder)
            return await asyncio.to_thread(lambda: wordDoc.WordDocFast(local_path, saveFolder).extractText())

        elif kind == "xlsx":
            local_path = await ensure_local_file(link, saveFolder)
            return await asyncio.to_thread(_extract_excel_text, local_path)

        elif kind in ("ppt", "pptx"):
            local_path = await ensure_local_file(link, saveFolder)
            return await asyncio.to_thread(_extract_ppt_text, local_path)

        elif kind == "zip":
            local_path = await ensure_local_file(link, saveFolder)
            return await asyncio.to_thread(_extract_zip_text, local_path, saveFolder)

        else:  # html — fetch and parse
            if not link.startswith("http") and "html" not in link:
                return ""
            html_content = await async_fetch_html(link)
            if extractHTML is None:
                return html_content
            html = extractHTML.HTML(htmlContent=html_content, htmlLink=link, htmlFile="")
            try:
                if hasattr(html, "async_getListSection"):
                    article_text = await html.async_getListSection()
                else:
                    article_text = await asyncio.to_thread(html.getListSection)
            except Exception:
                article_text = ""
            if not article_text:
                metadata_text = html.fetch_crossref_metadata(link)
                if metadata_text:
                    article_text = html.mergeTextInJson(metadata_text)
            return article_text

    except Exception as e:
        print(f"❌ async_extract_text failed for {link}: {e}")
        return ""
        

def extract_text(link,saveFolder):
  try:
      text = ""
      link = _github_blob_to_raw(link)
      name = _get_filename_from_url(link)
      print("name: ", name)
      local_temp_path = os.path.join(tempfile.gettempdir(), name)
      print("this is local temp path: ", local_temp_path)
      if os.path.exists(local_temp_path):
        input_to_class = local_temp_path
        print("exist")
      else:
        # 1. Check if file exists in shared Google Drive folder (only for real Drive folder IDs)
        _use_drive = pipeline is not None and not _is_local_path(saveFolder)
        file_id = pipeline.find_drive_file(name, saveFolder) if _use_drive else None
        if file_id:
            print("📥 Downloading from Google Drive...")
            pipeline.download_file_from_drive(name, saveFolder, local_temp_path)
        else:
            print("🌐 Downloading from web link...")
            try:
                content = _download_file(link)
                with open(local_temp_path, 'wb') as f:
                    f.write(content)
                print("✅ Saved locally.")
                # 2. Upload to Drive so it's available for later
                if _use_drive:
                    pipeline.upload_file_to_drive(local_temp_path, name, saveFolder)
            except Exception as e:
                print(f"❌ Download failed: {e}")
                return ""

        input_to_class = local_temp_path
        print(input_to_class)  
      kind = classify_url(link)
      if kind == "pdf" and pdf is not None:
        p = pdf.PDFFast(input_to_class, saveFolder)
        text = p.extract_text()
      elif kind == "docx" and wordDoc is not None:
        d = wordDoc.WordDocFast(input_to_class, saveFolder)
        text = d.extractText()
      elif kind == "xlsx":
        text = _extract_excel_text(input_to_class)
      elif kind in ("ppt", "pptx"):
        text = _extract_ppt_text(input_to_class)
      elif kind == "zip":
        text = _extract_zip_text(input_to_class, saveFolder)
      elif (link.startswith("http") or "html" in link) and extractHTML is not None:
        html = extractHTML.HTML("", link)
        text = html.getListSection()
      # Cleanup: delete the local temp file
      if name:
          if os.path.exists(local_temp_path):
            os.remove(local_temp_path)
            print(f"🧹 Deleted local temp file: {local_temp_path}")
      print("done extract text")
  except:
      text = ""
  return text

def extract_table(link,saveFolder):
  try:
      table = []
      link = _github_blob_to_raw(link)
      name = _get_filename_from_url(link)
      local_temp_path = os.path.join(tempfile.gettempdir(), name)
      if os.path.exists(local_temp_path):
        input_to_class = local_temp_path
        print("exist")
      else:
        # 1. Check if file exists in shared Google Drive folder (only for real Drive folder IDs)
        _use_drive = pipeline is not None and not _is_local_path(saveFolder)
        file_id = pipeline.find_drive_file(name, saveFolder) if _use_drive else None
        if file_id:
            print("📥 Downloading from Google Drive...")
            pipeline.download_file_from_drive(name, saveFolder, local_temp_path)
        else:
            print("🌐 Downloading from web link...")
            try:
                content = _download_file(link)
                with open(local_temp_path, 'wb') as f:
                    f.write(content)
                print("✅ Saved locally.")
                if _use_drive:
                    pipeline.upload_file_to_drive(local_temp_path, name, saveFolder)
            except Exception as e:
                print(f"❌ Download failed: {e}")
                return []

        input_to_class = local_temp_path
        print(input_to_class)
      kind = classify_url(link)
      if kind == "pdf" and pdf is not None:
        p = pdf.PDF(input_to_class, saveFolder)
        table = p.extractTable()
      elif kind == "docx" and wordDoc is not None:
        d = wordDoc.WordDocFast(input_to_class, saveFolder)
        table = d.extractTablesAsList()
      elif kind == "xlsx":
        try:
            xls = pd.ExcelFile(local_temp_path)
            table_list = []
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                table_list.append(df.fillna("").astype(str).values.tolist())
            table = table_list
        except Exception as e:
            print("❌ Failed to extract tables from Excel:", e)
      elif (link.startswith("http") or "html" in link) and extractHTML is not None:
        html = extractHTML.HTML("", link)
        table = html.extractTable()
      table = clean_tables_format(table)
      # Cleanup: delete the local temp file
      if os.path.exists(local_temp_path):
        os.remove(local_temp_path)
        print(f"🧹 Deleted local temp file: {local_temp_path}")
  except:
      table = []
  return table

def clean_tables_format(tables):
    """
    Ensures all tables are in consistent format: List[List[List[str]]]
    Cleans by:
    - Removing empty strings and rows
    - Converting all cells to strings
    - Handling DataFrames and list-of-lists
    """
    cleaned = []
    if tables:
      for table in tables:
          standardized = []

          # Case 1: Pandas DataFrame
          if isinstance(table, pd.DataFrame):
              table = table.fillna("").astype(str).values.tolist()

          # Case 2: List of Lists
          if isinstance(table, list) and all(isinstance(row, list) for row in table):
              for row in table:
                  filtered_row = [str(cell).strip() for cell in row if str(cell).strip()]
                  if filtered_row:
                      standardized.append(filtered_row)

          if standardized:
              cleaned.append(standardized)

    return cleaned

import json
def normalize_text_for_comparison(s: str) -> str:
    """
    Normalizes text for robust comparison by:
    1. Converting to lowercase.
    2. Replacing all types of newlines with a single consistent newline (\n).
    3. Removing extra spaces (e.g., multiple spaces, leading/trailing spaces on lines).
    4. Stripping leading/trailing whitespace from the entire string.
    """
    s = s.lower()
    s = s.replace('\r\n', '\n') # Handle Windows newlines
    s = s.replace('\r', '\n')   # Handle Mac classic newlines
    
    # Replace sequences of whitespace (including multiple newlines) with a single space
    # This might be too aggressive if you need to preserve paragraph breaks,
    # but good for exact word-sequence matching.
    s = re.sub(r'\s+', ' ', s) 
    
    return s.strip()
def merge_text_and_tables(text, tables, max_tokens=12000, keep_tables=True, tokenizer="cl100k_base", accession_id=None, isolate=None):
    """
    Merge cleaned text and table into one string for LLM input.
    - Avoids duplicating tables already in text
    - Extracts only relevant rows from large tables
    - Skips or saves oversized tables
    """
    import importlib
    json = importlib.import_module("json")

    def estimate_tokens(text_str):
        try:
            enc = tiktoken.get_encoding(tokenizer)
            return len(enc.encode(text_str))
        except:
            return len(text_str) // 4  # Fallback estimate

    def is_table_relevant(table, keywords, accession_id=None):
        flat = " ".join(" ".join(row).lower() for row in table)
        if accession_id and accession_id.lower() in flat:
            return True    
        return any(kw.lower() in flat for kw in keywords)
    preview, preview1 = "",""    
    llm_input = "## Document Text\n" + text.strip() + "\n"
    clean_text = normalize_text_for_comparison(text)

    if tables:
        for idx, table in enumerate(tables):
          keywords = ["province","district","region","village","location", "country", "region", "origin", "ancient", "modern"]
          if accession_id:  keywords += [accession_id.lower()]
          if isolate: keywords += [isolate.lower()]
          if is_table_relevant(table, keywords, accession_id):
            if len(table) > 0:
              for tab in table:
                preview = " ".join(tab) if tab else ""
                preview1 = "\n".join(tab) if tab else ""
                clean_preview = normalize_text_for_comparison(preview)
                clean_preview1 = normalize_text_for_comparison(preview1)
                if clean_preview not in clean_text:
                  if clean_preview1 not in clean_text:
                    table_str = json.dumps([tab], indent=2)
                    llm_input += f"## Table {idx+1}\n{table_str}\n"
    return llm_input.strip()

def preprocess_document(link, saveFolder, accession=None, isolate=None, article_text=None):
    if article_text:
      print("article text already available")
      text = article_text
    else:  
      try:
        print("start preprocess and extract text")
        text = extract_text(link, saveFolder)
      except: text = ""  
    try: 
      print("extract table start")
      success, the_output = pipeline.run_with_timeout(extract_table,args=(link,saveFolder),timeout=10)
      print("Returned from timeout logic")
      if success:
        tables = the_output#data_preprocess.merge_texts_skipping_overlap(all_output, final_input_link)
        print("yes succeed for extract table")
      else:
        print("not suceed etxract table")
        tables = []
      #tables = extract_table(link, saveFolder)
    except: tables = [] 
    if accession: accession = accession
    if isolate: isolate = isolate
    try:
      # print("merge text and table start")
      # success, the_output = pipeline.run_with_timeout(merge_text_and_tables,kwargs={"text":text,"tables":tables,"accession_id":accession, "isolate":isolate},timeout=30)
      # print("Returned from timeout logic")
      # if success:
      #   final_input = the_output#data_preprocess.merge_texts_skipping_overlap(all_output, final_input_link)
      #   print("yes succeed")
      # else:
      #   print("not suceed")
      print("just merge text and tables")
      final_input = text + ", ".join(tables)  
      #final_input = pipeline.timeout(merge_text_and_tables(text, tables, max_tokens=12000, accession_id=accession, isolate=isolate)
    except: 
      print("no succeed here in preprocess docu")
      final_input = ""
    return text, tables, final_input

def extract_sentences(text):
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return [s.strip() for s in sentences if s.strip()]

def is_irrelevant_number_sequence(text):
    if re.search(r'\b[A-Z]{2,}\d+\b|\b[A-Za-z]+\s+\d+\b', text, re.IGNORECASE):
        return False
    word_count = len(re.findall(r'\b[A-Za-z]{2,}\b', text))
    number_count = len(re.findall(r'\b\d[\d\.]*\b', text))
    total_tokens = len(re.findall(r'\S+', text))
    if total_tokens > 0 and (word_count / total_tokens < 0.2) and (number_count / total_tokens > 0.5):
        return True
    elif re.fullmatch(r'(\d+(\.\d+)?\s*)+', text.strip()):
        return True
    return False

def remove_isolated_single_digits(sentence):
    tokens = sentence.split()
    filtered_tokens = []
    for token in tokens:
        if token == '0' or token == '1':
            pass
        else:
            filtered_tokens.append(token)
    return ' '.join(filtered_tokens).strip()

def get_contextual_sentences_BFS(text_content, keyword, depth=2):
    def extract_codes(sentence):
    # Match codes like 'A1YU101', 'KM1', 'MO6' — at least 2 letters + numbers
      return [code for code in re.findall(r'\b[A-Z]{2,}[0-9]+\b', sentence, re.IGNORECASE)]
    sentences = extract_sentences(text_content)
    relevant_sentences = set()
    initial_keywords = set()

    # Define a regex to capture codes like A1YU101 or KM1
    # This pattern looks for an alphanumeric sequence followed by digits at the end of the string
    code_pattern = re.compile(r'([A-Z0-9]+?)(\d+)$', re.IGNORECASE)

    # Attempt to parse the keyword into its prefix and numerical part using re.search
    keyword_match = code_pattern.search(keyword)

    keyword_prefix = None
    keyword_num = None

    if keyword_match:
        keyword_prefix = keyword_match.group(1).lower()
        keyword_num = int(keyword_match.group(2))

    for sentence in sentences:
        sentence_added = False

        # 1. Check for exact match of the keyword
        if re.search(r'\b' + re.escape(keyword) + r'\b', sentence, re.IGNORECASE):
            relevant_sentences.add(sentence.strip())
            initial_keywords.add(keyword.lower())
            sentence_added = True

        # 2. Check for range patterns (e.g., A1YU101-A1YU137)
        # The range pattern should be broad enough to capture the full code string within the range.
        range_matches = re.finditer(r'([A-Z0-9]+-\d+)', sentence, re.IGNORECASE) # More specific range pattern if needed, or rely on full code pattern below
        range_matches = re.finditer(r'([A-Z0-9]+\d+)-([A-Z0-9]+\d+)', sentence, re.IGNORECASE) # This is the more robust range pattern

        for r_match in range_matches:
            start_code_str = r_match.group(1)
            end_code_str = r_match.group(2)

            # CRITICAL FIX: Use code_pattern.search for start_match and end_match
            start_match = code_pattern.search(start_code_str)
            end_match = code_pattern.search(end_code_str)

            if keyword_prefix and keyword_num is not None and start_match and end_match:
                start_prefix = start_match.group(1).lower()
                end_prefix = end_match.group(1).lower()
                start_num = int(start_match.group(2))
                end_num = int(end_match.group(2))

                # Check if the keyword's prefix matches and its number is within the range
                if keyword_prefix == start_prefix and \
                   keyword_prefix == end_prefix and \
                   start_num <= keyword_num <= end_num:
                    relevant_sentences.add(sentence.strip())
                    initial_keywords.add(start_code_str.lower())
                    initial_keywords.add(end_code_str.lower())
                    sentence_added = True
                    break # Only need to find one matching range per sentence

        # 3. If the sentence was added due to exact match or range, add all its alphanumeric codes
        #    to initial_keywords to ensure graph traversal from related terms.
        if sentence_added:
          for word in extract_codes(sentence):
            initial_keywords.add(word.lower())


    # Build word_to_sentences mapping for all sentences
    word_to_sentences = {}
    for sent in sentences:
      codes_in_sent = set(extract_codes(sent))
      for code in codes_in_sent:
          word_to_sentences.setdefault(code.lower(), set()).add(sent.strip())


    # Build the graph
    graph = {}
    for sent in sentences:
      codes = set(extract_codes(sent))
      for word1 in codes:
          word1_lower = word1.lower()
          graph.setdefault(word1_lower, set())
          for word2 in codes:
              word2_lower = word2.lower()
              if word1_lower != word2_lower:
                  graph[word1_lower].add(word2_lower)


    # Perform BFS/graph traversal
    queue = [(k, 0) for k in initial_keywords if k in word_to_sentences]
    visited_words = set(initial_keywords)

    while queue:
        current_word, level = queue.pop(0)
        if level >= depth:
            continue

        relevant_sentences.update(word_to_sentences.get(current_word, []))

        for neighbor in graph.get(current_word, []):
            if neighbor not in visited_words:
                visited_words.add(neighbor)
                queue.append((neighbor, level + 1))

    final_sentences = set()
    for sentence in relevant_sentences:
        if not is_irrelevant_number_sequence(sentence):
            processed_sentence = remove_isolated_single_digits(sentence)
            if processed_sentence:
                final_sentences.add(processed_sentence)

    return "\n".join(sorted(list(final_sentences)))



def get_contextual_sentences_DFS(text_content, keyword, depth=2):
    sentences = extract_sentences(text_content)

    # Build word-to-sentences mapping
    word_to_sentences = {}
    for sent in sentences:
        words_in_sent = set(re.findall(r'\b[A-Za-z0-9\-_\/]+\b', sent))
        for word in words_in_sent:
            word_to_sentences.setdefault(word.lower(), set()).add(sent.strip())

    # Function to extract codes in a sentence
    def extract_codes(sentence):
      # Only codes like 'KSK1', 'MG272794', not pure numbers
      return [code for code in re.findall(r'\b[A-Z]{2,}[0-9]+\b', sentence, re.IGNORECASE)]

    # DFS with priority based on distance to keyword and early stop if country found
    def dfs_traverse(current_word, current_depth, max_depth, visited_words, collected_sentences, parent_sentence=None):
        country = "unknown"
        if current_depth > max_depth:
            return country, False

        if current_word not in word_to_sentences:
            return country, False

        for sentence in word_to_sentences[current_word]:
            if sentence == parent_sentence:
                continue  # avoid reusing the same sentence

            collected_sentences.add(sentence)

            #print("current_word:", current_word)
            small_sen = extract_context(sentence, current_word, int(len(sentence) / 4))
            #print(small_sen)
            country = model.get_country_from_text(small_sen)
            #print("small context country:", country)
            if country.lower() != "unknown":
                return country, True
            else:
                country = model.get_country_from_text(sentence)
                #print("full sentence country:", country)
                if country.lower() != "unknown":
                    return country, True

            codes_in_sentence = extract_codes(sentence)
            idx = next((i for i, code in enumerate(codes_in_sentence) if code.lower() == current_word.lower()), None)
            if idx is None:
                continue

            sorted_children = sorted(
                [code for code in codes_in_sentence if code.lower() not in visited_words],
                key=lambda x: (abs(codes_in_sentence.index(x) - idx),
                               0 if codes_in_sentence.index(x) > idx else 1)
            )

            #print("sorted_children:", sorted_children)
            for child in sorted_children:
                child_lower = child.lower()
                if child_lower not in visited_words:
                    visited_words.add(child_lower)
                    country, should_stop = dfs_traverse(
                        child_lower, current_depth + 1, max_depth,
                        visited_words, collected_sentences, parent_sentence=sentence
                    )
                    if should_stop:
                        return country, True

        return country, False

    # Begin DFS
    collected_sentences = set()
    visited_words = set([keyword.lower()])
    country, status = dfs_traverse(keyword.lower(), 0, depth, visited_words, collected_sentences)

    # Filter irrelevant sentences
    final_sentences = set()
    for sentence in collected_sentences:
        if not is_irrelevant_number_sequence(sentence):
            processed = remove_isolated_single_digits(sentence)
            if processed:
                final_sentences.add(processed)
    if not final_sentences:
      return country, text_content
    return country, "\n".join(sorted(list(final_sentences)))

# Helper function for normalizing text for overlap comparison
def normalize_for_overlap(s: str) -> str:
    s = re.sub(r'[^a-zA-Z0-9\s]', ' ', s).lower()
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def merge_texts_skipping_overlap(text1: str, text2: str) -> str:
    if not text1: return text2
    if not text2: return text1

    # Case 1: text2 is fully contained in text1 or vice-versa
    if text2 in text1:
        return text1
    if text1 in text2:
        return text2

    # --- Option 1: Original behavior (suffix of text1, prefix of text2) ---
    # This is what your function was primarily designed for.
    # It looks for the overlap at the "junction" of text1 and text2.
    
    max_junction_overlap = 0
    for i in range(min(len(text1), len(text2)), 0, -1):
        suffix1 = text1[-i:]
        prefix2 = text2[:i]
        # Prioritize exact match, then normalized match
        if suffix1 == prefix2:
            max_junction_overlap = i
            break
        elif normalize_for_overlap(suffix1) == normalize_for_overlap(prefix2):
            max_junction_overlap = i
            break # Take the first (longest) normalized match

    if max_junction_overlap > 0:
        merged_text = text1 + text2[max_junction_overlap:]
        return re.sub(r'\s+', ' ', merged_text).strip()

    # --- Option 2: Longest Common Prefix (for cases like "Hi, I am Vy.") ---
    # This addresses your specific test case where the overlap is at the very beginning of both strings.
    # This is often used when trying to deduplicate content that shares a common start.

    longest_common_prefix_len = 0
    min_len = min(len(text1), len(text2))
    for i in range(min_len):
        if text1[i] == text2[i]:
            longest_common_prefix_len = i + 1
        else:
            break
    
    # If a common prefix is found AND it's a significant portion (e.g., more than a few chars)
    # AND the remaining parts are distinct, then apply this merge.
    # This is a heuristic and might need fine-tuning.
    if longest_common_prefix_len > 0 and \
       text1[longest_common_prefix_len:].strip() and \
       text2[longest_common_prefix_len:].strip():

        # Only merge this way if the remaining parts are not empty (i.e., not exact duplicates)
        # For "Hi, I am Vy. Nice to meet you." and "Hi, I am Vy. Goodbye Vy."
        # common prefix is "Hi, I am Vy."
        # Remaining text1: " Nice to meet you."
        # Remaining text2: " Goodbye Vy."
        # So we merge common_prefix + remaining_text1 + remaining_text2
        
        common_prefix_str = text1[:longest_common_prefix_len]
        remainder_text1 = text1[longest_common_prefix_len:]
        remainder_text2 = text2[longest_common_prefix_len:]
        
        merged_text = common_prefix_str + remainder_text1 + remainder_text2
        return re.sub(r'\s+', ' ', merged_text).strip()


    # If neither specific overlap type is found, just concatenate
    merged_text = text1 + text2
    return re.sub(r'\s+', ' ', merged_text).strip()

from docx import Document
from pipeline import upload_file_to_drive    
# def save_text_to_docx(text_content: str, file_path: str):
#     """
#     Saves a given text string into a .docx file.

#     Args:
#         text_content (str): The text string to save.
#         file_path (str): The full path including the filename where the .docx file will be saved.
#                          Example: '/content/drive/MyDrive/CollectData/Examples/test/SEA_1234/merged_document.docx'
#     """
#     try:
#         document = Document()

#         # Add the entire text as a single paragraph, or split by newlines for multiple paragraphs
#         for paragraph_text in text_content.split('\n'):
#             document.add_paragraph(paragraph_text)

#         document.save(file_path)
#         print(f"Text successfully saved to '{file_path}'")
#     except Exception as e:
#         print(f"Error saving text to docx file: {e}") 
# def save_text_to_docx(text_content: str, filename: str, drive_folder_id: str):
#     """
#     Saves a given text string into a .docx file locally, then uploads to Google Drive.

#     Args:
#         text_content (str): The text string to save.
#         filename (str): The target .docx file name, e.g. 'BRU18_merged_document.docx'.
#         drive_folder_id (str): Google Drive folder ID where to upload the file.
#     """
#     try:
#         # ✅ Save to temporary local path first
#         print("file name: ", filename)
#         print("length text content: ", len(text_content))
#         local_path = os.path.join(tempfile.gettempdir(), filename)
#         document = Document()
#         for paragraph_text in text_content.split('\n'):
#             document.add_paragraph(paragraph_text)
#         document.save(local_path)
#         print(f"✅ Text saved locally to: {local_path}")

#         # ✅ Upload to Drive
#         pipeline.upload_file_to_drive(local_path, filename, drive_folder_id)
#         print(f"✅ Uploaded '{filename}' to Google Drive folder ID: {drive_folder_id}")

#     except Exception as e:
#         print(f"❌ Error saving or uploading DOCX: {e}")
def save_text_to_docx(text_content: str, full_local_path: str):
    document = Document()
    for paragraph_text in text_content.split('\n'):
        document.add_paragraph(paragraph_text)
    document.save(full_local_path)
    print(f"✅ Saved DOCX locally: {full_local_path}")



'''2 scenerios:
- quick look then found then deepdive and directly get location then stop
- quick look then found then deepdive but not find location then hold the related words then 
look another files iteratively for each related word and find location and stop'''
def extract_context(text, keyword, window=500):
    # firstly try accession number
    code_pattern = re.compile(r'([A-Z0-9]+?)(\d+)$', re.IGNORECASE)

    # Attempt to parse the keyword into its prefix and numerical part using re.search
    keyword_match = code_pattern.search(keyword)

    keyword_prefix = None
    keyword_num = None

    if keyword_match:
        keyword_prefix = keyword_match.group(1).lower()
        keyword_num = int(keyword_match.group(2))
    text = text.lower()    
    idx = text.find(keyword.lower())
    if idx == -1:
      if keyword_prefix:
        idx = text.find(keyword_prefix)
      if idx == -1:
        return "Sample ID not found."
      return text[max(0, idx-window): idx+window]  
    return text[max(0, idx-window): idx+window]
def process_inputToken(filePaths, saveLinkFolder,accession=None, isolate=None):
  cache = {}
  country = "unknown"
  output = ""
  tem_output, small_output = "",""
  keyword_appear = (False,"")
  keywords = []
  if isolate: keywords.append(isolate)
  if accession: keywords.append(accession)
  for f in filePaths:
    # scenerio 1: direct location: truncate the context and then use qa model?
    if keywords:
      for keyword in keywords:
        text, tables, final_input = preprocess_document(f,saveLinkFolder, isolate=keyword)
        if keyword in final_input:
          context = extract_context(final_input, keyword)
          # quick look if country already in context and if yes then return
          country = model.get_country_from_text(context)
          if country != "unknown":
            return country, context, final_input
          else:
            country = model.get_country_from_text(final_input)  
            if country != "unknown":
              return country, context, final_input
            else: # might be cross-ref
              keyword_appear = (True, f)
              cache[f] = context
              small_output = merge_texts_skipping_overlap(output, context) + "\n"
              chunkBFS = get_contextual_sentences_BFS(small_output, keyword)
              countryBFS = model.get_country_from_text(chunkBFS)
              countryDFS, chunkDFS = get_contextual_sentences_DFS(output, keyword)
              output = merge_texts_skipping_overlap(output, final_input)
              if countryDFS != "unknown" and countryBFS != "unknown":
                if len(chunkDFS) <= len(chunkBFS):
                  return countryDFS, chunkDFS, output
                else:
                  return countryBFS, chunkBFS, output
              else:        
                if countryDFS != "unknown":  
                  return countryDFS, chunkDFS, output
                if countryBFS != "unknown":
                  return countryBFS, chunkBFS, output
        else:
        # scenerio 2: 
          '''cross-ref: ex: A1YU101 keyword in file 2 which includes KM1 but KM1 in file 1 
          but if we look at file 1 first then maybe we can have lookup dict which country 
          such as Thailand as the key and its re''' 
          cache[f] = final_input
          if keyword_appear[0] == True:
            for c in cache:
              if c!=keyword_appear[1]:
                if cache[c].lower() not in output.lower():
                  output = merge_texts_skipping_overlap(output, cache[c]) + "\n"
                  chunkBFS = get_contextual_sentences_BFS(output, keyword)
                  countryBFS = model.get_country_from_text(chunkBFS)
                  countryDFS, chunkDFS = get_contextual_sentences_DFS(output, keyword)
                  if countryDFS != "unknown" and countryBFS != "unknown":
                    if len(chunkDFS) <= len(chunkBFS):
                      return countryDFS, chunkDFS, output
                    else:
                      return countryBFS, chunkBFS, output
                  else:        
                    if countryDFS != "unknown":  
                      return countryDFS, chunkDFS, output
                    if countryBFS != "unknown":
                      return countryBFS, chunkBFS, output
          else:
            if cache[f].lower() not in output.lower():
              output = merge_texts_skipping_overlap(output, cache[f]) + "\n"          
  if len(output) == 0 or keyword_appear[0]==False:
    for c in cache:
      if cache[c].lower() not in output.lower():
        output = merge_texts_skipping_overlap(output, cache[c]) + "\n"
  return country, "", output   

import re
from typing import List, Tuple
from collections import defaultdict
try:
    import tiktoken
    enc = tiktoken.get_encoding("cl100k_base")
except Exception:
    tiktoken = None
    enc = None
# =============================
# 1. HELPER: Token counter
# =============================
def num_tokens(text: str) -> int:
    if enc is None:
        return len(text) // 4  # rough estimate: ~4 chars per token
    return len(enc.encode(text))

# =============================
# 2. Extract context window
# =============================
def extract_context_reduceToken(text: str, keyword: str, window: int = 1200) -> str:
    pattern = re.escape(keyword)
    match = re.search(pattern, text, flags=re.IGNORECASE)
    if not match:
        return ""
    start = max(0, match.start() - window)
    end = match.end() + window
    return text[start:end]

# =============================
# 3. Expand question keywords dynamically
# =============================
try:
    from sentence_transformers import SentenceTransformer, util as _st_util
    _sentence_model = SentenceTransformer("paraphrase-MiniLM-L6-v2")
except Exception:
    SentenceTransformer = _st_util = _sentence_model = None

import nltk
nltk.download('wordnet')
import re
from nltk.corpus import wordnet as wn
from nltk.corpus import stopwords

STOPWORDS = set(stopwords.words("english"))

def synonym_expand(word: str):
    syns = set()
    for synset in wn.synsets(word):
        for lemma in synset.lemmas():
            syns.add(lemma.name().replace("_", " "))
    return list(syns)

def normalize_keyword(kw: str):
    """
    Convert things like:
    - 'country_name' → ['country', 'name']
    - 'sample-type' → ['sample', 'type']
    - 'host_species' → ['host', 'species']
    
    Remove stopwords like 'name', 'type', 'data'.
    """
    #parts = re.split(r"[_\-\s]+", kw.lower())
    #return [p for p in parts if p and p not in STOPWORDS]
    parts = re.split(r"[\/,;_\-\s]+", kw.strip().lower())

    return [p for p in parts if p and p != "unknown"]

def expand_keywords(accession: str, question_kws: list) -> list:
    dynamic = [accession]

    for kw in question_kws:
        # Normalize multiword keyword
        base_words = normalize_keyword(kw)
        dynamic.extend(base_words)

        # Expand each base word
        for w in base_words:
            dynamic.extend(synonym_expand(w))

    # remove duplicates and empty values
    dynamic = [w for w in set(dynamic) if w.strip()]
    return dynamic

# =============================
# 4. RELEVANCE SCORING
# =============================
def score_context(context: str, keywords: List[str]) -> int:
    score = 0
    lowered = context.lower()
    for kw in keywords:
        if kw.lower() in lowered:
            score += 1
    return score

# =============================
# 5. MAIN FUNCTION
# =============================
def build_context_for_llm(texts: List[str], accession: str, question_kws: List[str], max_tokens: int = 1_000_000):
    keywords = expand_keywords(accession, question_kws)

    primary_contexts = []
    supplemental_contexts = []
    primary_found = False

    # Step 1: Extract contexts for all keywords from all texts
    for text in texts:
        text_l = text.lower()
        print("this is len text: ", len(text_l))
        # PRIORITY 1: Accession context
        if accession.lower() in text_l:
            ctx = extract_context_reduceToken(text, accession, window=3000)
            if ctx:
                print("acc in context")
                primary_contexts.append(ctx)
                primary_found = True

        # PRIORITY 2: keyword contexts
        for kw in keywords:
            ctx = extract_context_reduceToken(text, kw, window=1200)
            if ctx:
                supplemental_contexts.append(ctx)

    # Step 2: Score all supplemental contexts
    scored = [(ctx, score_context(ctx, keywords)) for ctx in supplemental_contexts]
    scored.sort(key=lambda x: -x[1])

    # Step 3: Merge primary contexts first
    final_context = "\n\n--- PRIMARY CONTEXT ---\n\n" + "\n\n---\n\n".join(primary_contexts)

    # Step 4: Add supplemental contexts until token budget is full
    for ctx, score in scored:
        if num_tokens(final_context) >= max_tokens:
            break
        final_context += f"\n\n--- RELEVANT CONTEXT (score={score}) ---\n\n{ctx}"

    # Final safety trim
    if num_tokens(final_context) > max_tokens:
        final_context = final_context[:300000]  # crude but safe cutoff

    return final_context