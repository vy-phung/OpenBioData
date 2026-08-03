import re
import json
import csv
import os
import asyncio
from collections import defaultdict
import ast
import math
try:
    import pycountry
except ImportError:
    pycountry = None
try:
    from docx import Document
except ImportError:
    Document = None
try:
    import numpy as np
    import faiss
except ImportError:
    np = faiss = None
try:
    import data_preprocess
except ImportError:
    data_preprocess = None
try:
    import mtdna_classifier
except ImportError:
    mtdna_classifier = None
try:
    import smart_fallback
except ImportError:
    smart_fallback = None
try:
    import pipeline
except ImportError:
    pipeline = None
# --- IMPORTANT: UNCOMMENT AND CONFIGURE YOUR REAL API KEY ---
import google.generativeai as genai

#genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
#genai.configure(api_key=os.getenv("GOOGLE_API_KEY_BACKUP"))
genai.configure(api_key=os.getenv("NEW_GOOGLE_API_KEY") or os.getenv("GOOGLE_API_KEY") or os.getenv("NEW_GEMINI_API"))

import nltk
from nltk.corpus import stopwords
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')
nltk.download('punkt_tab')    
# # --- Define Pricing Constants (for Gemini 1.5 Flash & text-embedding-004) ---
# # Prices are per 1,000 tokens
# PRICE_PER_1K_INPUT_LLM = 0.000075  # $0.075 per 1M tokens
# PRICE_PER_1K_OUTPUT_LLM = 0.0003   # $0.30 per 1M tokens
# PRICE_PER_1K_EMBEDDING_INPUT = 0.000025 # $0.025 per 1M tokens

# Gemini 2.5 Flash-Lite pricing per 1,000 tokens
PRICE_PER_1K_INPUT_LLM = 0.00010      # $0.10 per 1M input tokens
PRICE_PER_1K_OUTPUT_LLM = 0.00040     # $0.40 per 1M output tokens

# Embedding-001 pricing per 1,000 input tokens
PRICE_PER_1K_EMBEDDING_INPUT = 0.00015  # $0.15 per 1M input tokens
# --- API Functions (REAL API FUNCTIONS) ---

# def get_embedding(text, task_type="RETRIEVAL_DOCUMENT"):
#     """Generates an embedding for the given text using a Google embedding model."""
#     try:
#         result = genai.embed_content(
#             model="models/text-embedding-004", # Specify the embedding model
#             content=text,
#             task_type=task_type
#         )
#         return np.array(result['embedding']).astype('float32')
#     except Exception as e:
#         print(f"Error getting embedding: {e}")
#         return np.zeros(768, dtype='float32')
def get_embedding(text, task_type="RETRIEVAL_DOCUMENT"):
    """Safe Gemini 1.5 embedding call with fallback."""
    import numpy as np
    try:
        if not text or len(text.strip()) == 0:
            raise ValueError("Empty text cannot be embedded.")
        result = genai.embed_content(
            model="models/text-embedding-004",
            content=text,
            task_type=task_type
        )
        return np.array(result['embedding'], dtype='float32')
    except Exception as e:
        print(f"❌ Embedding error: {e}")
        return np.zeros(768, dtype='float32')


_HAIKU_MODEL = "claude-haiku-4-5-20251001"
_SONNET_MODEL = "claude-sonnet-5"
_ANTHROPIC_MAX_TOKENS = 8192  # output budget; also used as the "expected output" reserve below

# Route to Haiku (200K hard context limit) while comfortably clear of it, to
# Sonnet 5 (1M context) once Haiku's margin is used up, and only fall through
# to Gemini if the context is too large even for Sonnet or the Anthropic call
# itself errors -- see call_llm_api()'s model-selection block below.
_HAIKU_SAFE_INPUT_BUDGET = 170_000    # real margin below Haiku's 200K hard limit
_SONNET_CONTEXT_LIMIT = 1_000_000     # Sonnet 5's context window


def _estimate_tokens(text: str) -> int:
    """Rough token estimate for routing only (not billing): chars / 4."""
    return len(text or "") // 4


from dev_llm_cache import dev_cache_wrap


@dev_cache_wrap
def call_llm_api(prompt, model_name=None):
    """Call LLM — tries Anthropic first (routing between Haiku and Sonnet 5
    by estimated context size; see model-selection block below), then each
    Gemini key in order as a true last resort.
    Set env var SKIP_LLM_API=true to skip all LLM calls (returns 'unknown' placeholders
    for testing NCBI resolution without consuming API credits).
    """
    if os.getenv('SKIP_LLM_API', '').lower() in ('1', 'true', 'yes'):
        print('[SKIP_LLM_API] LLM call skipped (test mode) — returning unknown placeholders')
        return 'unknown, unknown', None

    last_error = None

    # --- 1. Anthropic (ANTHROPIC_API_KEY) ---
    anthropic_key = os.getenv("ANTHROPIC_API_KEY")
    if anthropic_key:
        estimated_input_tokens = _estimate_tokens(prompt)
        # Sonnet 5 rejects a manual budget_tokens thinking config (400) and any
        # non-default temperature/top_p/top_k (400) -- deliberately NOT setting
        # those here. Thinking itself is explicitly disabled below for Sonnet.
        if estimated_input_tokens + _ANTHROPIC_MAX_TOKENS <= _HAIKU_SAFE_INPUT_BUDGET:
            chosen_model = _HAIKU_MODEL
        elif estimated_input_tokens + _ANTHROPIC_MAX_TOKENS <= _SONNET_CONTEXT_LIMIT:
            chosen_model = _SONNET_MODEL
        else:
            chosen_model = None  # too large even for Sonnet 5 -- skip Anthropic, try Gemini

        if chosen_model:
            print(f"[call_llm_api] routing: model={chosen_model} estimated_input_tokens={estimated_input_tokens}")
            try:
                import anthropic as _anthropic
                client = _anthropic.Anthropic(api_key=anthropic_key)
                create_kwargs = dict(
                    model=chosen_model,
                    max_tokens=_ANTHROPIC_MAX_TOKENS,
                    messages=[{"role": "user", "content": prompt}],
                )
                # Sonnet 5 runs adaptive thinking by default when `thinking` is
                # omitted, unlike Haiku. Every call_llm_api call site in this
                # codebase is structured field extraction from provided text
                # (never open-ended multi-step reasoning), so thinking buys
                # little here while eating into max_tokens -- it has caused
                # Sonnet to spend the whole budget thinking and return a
                # response with a ThinkingBlock but no text block at all.
                # Disabling it removes that failure mode outright (Sonnet 5
                # accepts thinking: disabled; Haiku has no thinking to disable).
                if chosen_model == _SONNET_MODEL:
                    create_kwargs["thinking"] = {"type": "disabled"}
                msg = client.messages.create(**create_kwargs)
                text_block = next((b for b in msg.content if getattr(b, "type", None) == "text"), None)
                if text_block is None:
                    # Belt-and-suspenders: thinking is disabled above, so this
                    # shouldn't recur, but if Anthropic ever returns a
                    # text-less response for another reason, retry once on
                    # Claude itself with double the output budget before
                    # giving up -- never let this fall through to the Gemini
                    # branch below, which would silently swap providers for a
                    # response Anthropic never actually finished producing.
                    print(f"[call_llm_api] {chosen_model} returned no text block "
                          f"(block types: {[getattr(b, 'type', None) for b in msg.content]}) "
                          f"-- retrying on {chosen_model} with max_tokens={_ANTHROPIC_MAX_TOKENS * 2}")
                    retry_kwargs = dict(create_kwargs, max_tokens=_ANTHROPIC_MAX_TOKENS * 2)
                    msg = client.messages.create(**retry_kwargs)
                    text_block = next((b for b in msg.content if getattr(b, "type", None) == "text"), None)
                    if text_block is None:
                        raise RuntimeError(
                            f"Anthropic response for {chosen_model} had no text content block "
                            f"after retry with doubled max_tokens "
                            f"(block types: {[getattr(b, 'type', None) for b in msg.content]})"
                        )
                print(f"[call_llm_api] used: model={chosen_model} estimated_input_tokens={estimated_input_tokens}")
                return text_block.text, None
            except Exception as e:
                last_error = e
                err_str = str(e).lower()
                if "429" in str(e) or "rate_limit" in err_str or "overloaded" in err_str:
                    raise  # let safe_call_llm retry
                if "no text content block" in err_str:
                    raise  # exhausted retries on Claude itself -- do not fall through to Gemini
                print(f"Anthropic API error ({chosen_model}): {e} — trying Gemini keys.")
        else:
            print(f"[call_llm_api] routing: estimated_input_tokens={estimated_input_tokens} exceeds "
                  f"Sonnet 5's context window — skipping Anthropic, trying Gemini.")

    # --- 2. Gemini — try each key in order ---
    gemini_model = model_name or "gemini-2.5-flash-lite"
    gemini_keys = [
        ("NEW_GOOGLE_API_KEY", os.getenv("NEW_GOOGLE_API_KEY")),
        ("GOOGLE_API_KEY",     os.getenv("GOOGLE_API_KEY")),
        ("NEW_GEMINI_API",     os.getenv("NEW_GEMINI_API")),
    ]
    for key_name, key_val in gemini_keys:
        if not key_val:
            continue
        try:
            genai.configure(api_key=key_val)
            m = genai.GenerativeModel(gemini_model)
            response = m.generate_content(prompt)
            return response.text, m
        except Exception as e:
            last_error = e
            err_str = str(e).lower()
            if "429" in str(e) or "rate" in err_str:
                raise  # rate limit — let safe_call_llm retry with backoff
            print(f"Gemini error with {key_name}: {e} — trying next key.")

    raise RuntimeError(f"All LLM API keys exhausted. Last error: {last_error}")


# --- Core Document Processing Functions (All previously provided and fixed) ---

def read_docx_text(path):
    """
    Reads text and extracts potential table-like strings from a .docx document.
    Separates plain text from structured [ [ ] ] list-like tables.
    Also attempts to extract a document title.
    """
    doc = Document(path)
    plain_text_paragraphs = []
    table_strings = []
    document_title = "Unknown Document Title" # Default

    # Attempt to extract the document title from the first few paragraphs
    title_paragraphs = [p.text.strip() for p in doc.paragraphs[:5] if p.text.strip()]
    if title_paragraphs:
        # A heuristic to find a title: often the first or second non-empty paragraph
        # or a very long first paragraph if it's the title
        if len(title_paragraphs[0]) > 50 and "Human Genetics" not in title_paragraphs[0]:
            document_title = title_paragraphs[0]
        elif len(title_paragraphs) > 1 and len(title_paragraphs[1]) > 50 and "Human Genetics" not in title_paragraphs[1]:
            document_title = title_paragraphs[1]
        elif any("Complete mitochondrial genomes" in p for p in title_paragraphs):
            # Fallback to a known title phrase if present
            document_title = "Complete mitochondrial genomes of Thai and Lao populations indicate an ancient origin of Austroasiatic groups and demic diffusion in the spread of Tai–Kadai languages"

    current_table_lines = []
    in_table_parsing_mode = False

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # Condition to start or continue table parsing
        if text.startswith("## Table "): # Start of a new table section
            if in_table_parsing_mode and current_table_lines:
                table_strings.append("\n".join(current_table_lines))
            current_table_lines = [text] # Include the "## Table X" line
            in_table_parsing_mode = True
        elif in_table_parsing_mode and (text.startswith("[") or text.startswith('"')):
            # Continue collecting lines if we're in table mode and it looks like table data
            # Table data often starts with '[' for lists, or '"' for quoted strings within lists.
            current_table_lines.append(text)
        else:
            # If not in table mode, or if a line doesn't look like table data,
            # then close the current table (if any) and add the line to plain text.
            if in_table_parsing_mode and current_table_lines:
                table_strings.append("\n".join(current_table_lines))
                current_table_lines = []
            in_table_parsing_mode = False
            plain_text_paragraphs.append(text)

    # After the loop, add any remaining table lines
    if current_table_lines:
        table_strings.append("\n".join(current_table_lines))

    return "\n".join(plain_text_paragraphs), table_strings, document_title

# --- Structured Data Extraction and RAG Functions ---

def parse_literal_python_list(table_str):
    list_match = re.search(r'(\[\s*\[\s*(?:.|\n)*?\s*\]\s*\])', table_str)
    #print("Debug: list_match object (before if check):", list_match)
    if not list_match:
        if "table" in table_str.lower(): # then the table doest have the "]]" at the end
            table_str += "]]"
            list_match = re.search(r'(\[\s*\[\s*(?:.|\n)*?\s*\]\s*\])', table_str)
    if list_match:
        try:
            matched_string = list_match.group(1)
            #print("Debug: Matched string for literal_eval:", matched_string)
            return ast.literal_eval(matched_string)
        except (ValueError, SyntaxError) as e:
            print(f"Error evaluating literal: {e}")
            return []
    return []


_individual_code_parser = re.compile(r'([A-Z0-9]+?)(\d+)$', re.IGNORECASE)
def _parse_individual_code_parts(code_str):
    match = _individual_code_parser.search(code_str)
    if match:
        return match.group(1), match.group(2)
    return None, None


def parse_sample_id_to_population_code(plain_text_content):
    sample_id_map = {}
    contiguous_ranges_data = defaultdict(list)

    #section_start_marker = "The sample identification of each population is as follows:"
    section_start_marker = ["The sample identification of each population is as follows:","## table"]
    
    for s in section_start_marker:
      relevant_text_search = re.search(
          re.escape(s.lower()) + r"\s*(.*?)(?=\n##|\Z)",
          plain_text_content.lower(),
          re.DOTALL
      )
      if relevant_text_search: 
        break
      
    if not relevant_text_search:
        print("Warning: 'Sample ID Population Code' section start marker not found or block empty.")
        return sample_id_map, contiguous_ranges_data

    relevant_text_block = relevant_text_search.group(1).strip()

    # print(f"\nDEBUG_PARSING: --- Start of relevant_text_block (first 500 chars) ---")
    # print(relevant_text_block[:500])
    # print(f"DEBUG_PARSING: --- End of relevant_text_block (last 500 chars) ---")
    # print(relevant_text_block[-500:])
    # print(f"DEBUG_PARSING: Relevant text block length: {len(relevant_text_block)}")

    mapping_pattern = re.compile(
    r'\b([A-Z0-9]+\d+)(?:-([A-Z0-9]+\d+))?\s+([A-Z0-9]+)\b', # Changed the last group
    re.IGNORECASE)

    range_expansion_count = 0
    direct_id_count = 0
    total_matches_found = 0
    for match in mapping_pattern.finditer(relevant_text_block):
        total_matches_found += 1
        id1_full_str, id2_full_str_opt, pop_code = match.groups()

        #print(f"  DEBUG_PARSING: Matched: '{match.group(0)}'")

        pop_code_upper = pop_code.upper()

        id1_prefix, id1_num_str = _parse_individual_code_parts(id1_full_str)
        if id1_prefix is None:
            #print(f"    DEBUG_PARSING: Failed to parse ID1: {id1_full_str}. Skipping this mapping.")
            continue

        if id2_full_str_opt:
            id2_prefix_opt, id2_num_str_opt = _parse_individual_code_parts(id2_full_str_opt)
            if id2_prefix_opt is None:
                #print(f"    DEBUG_PARSING: Failed to parse ID2: {id2_full_str_opt}. Treating {id1_full_str} as single ID1.")
                sample_id_map[f"{id1_prefix.upper()}{id1_num_str}"] = pop_code_upper
                direct_id_count += 1
                continue

            #print(f"    DEBUG_PARSING: Comparing prefixes: '{id1_prefix.lower()}' vs '{id2_prefix_opt.lower()}'")
            if id1_prefix.lower() == id2_prefix_opt.lower():
                #print(f"    DEBUG_PARSING: ---> Prefixes MATCH for range expansion! Range: {id1_prefix}{id1_num_str}-{id2_prefix_opt}{id2_num_str_opt}")
                try:
                    start_num = int(id1_num_str)
                    end_num = int(id2_num_str_opt)
                    for num in range(start_num, end_num + 1):
                        sample_id = f"{id1_prefix.upper()}{num}"
                        sample_id_map[sample_id] = pop_code_upper
                        range_expansion_count += 1
                    contiguous_ranges_data[id1_prefix.upper()].append(
                        (start_num, end_num, pop_code_upper)
                    )
                except ValueError:
                    print(f"        DEBUG_PARSING: ValueError in range conversion for {id1_num_str}-{id2_num_str_opt}. Adding endpoints only.")
                    sample_id_map[f"{id1_prefix.upper()}{id1_num_str}"] = pop_code_upper
                    sample_id_map[f"{id2_prefix_opt.upper()}{id2_num_str_opt}"] = pop_code_upper
                    direct_id_count += 2
            else:
                #print(f"    DEBUG_PARSING: Prefixes MISMATCH for range: '{id1_prefix}' vs '{id2_prefix_opt}'. Adding endpoints only.")
                sample_id_map[f"{id1_prefix.upper()}{id1_num_str}"] = pop_code_upper
                sample_id_map[f"{id2_prefix_opt.upper()}{id2_num_str_opt}"] = pop_code_upper
                direct_id_count += 2
        else:
            sample_id_map[f"{id1_prefix.upper()}{id1_num_str}"] = pop_code_upper
            direct_id_count += 1

    # print(f"DEBUG_PARSING: Total matches found by regex: {total_matches_found}.")
    # print(f"DEBUG_PARSING: Parsed sample IDs: {len(sample_id_map)} total entries.")
    # print(f"DEBUG_PARSING:   (including {range_expansion_count} from range expansion and {direct_id_count} direct ID/endpoint entries).")
    return sample_id_map, contiguous_ranges_data

country_keywords_regional_overrides = {
    "north thailand": "Thailand", "central thailand": "Thailand",
    "northeast thailand": "Thailand", "east myanmar": "Myanmar", "west thailand": "Thailand",
    "central india": "India", "east india": "India", "northeast india": "India",
    "south sibera": "Russia", "siberia": "Russia", "yunnan": "China", #"tibet": "China",
    "sumatra": "Indonesia", "borneo": "Indonesia",
    "northern mindanao": "Philippines", "west malaysia": "Malaysia",
    "mongolia": "China",
    "beijing": "China",
    "north laos": "Laos", "central laos": "Laos",
    "east myanmar": "Myanmar", "west myanmar": "Myanmar"}

# Updated get_country_from_text function
def get_country_from_text(text):
    text_lower = text.lower()

    # 1. Use pycountry for official country names and common aliases
    for country in pycountry.countries:
        # Check full name match first
        if text_lower == country.name.lower():
            return country.name
        
        # Safely check for common_name
        if hasattr(country, 'common_name') and text_lower == country.common_name.lower():
            return country.common_name
            
        # Safely check for official_name
        if hasattr(country, 'official_name') and text_lower == country.official_name.lower():
            return country.official_name

        # Check if country name is part of the text (e.g., 'Thailand' in 'Thailand border')
        if country.name.lower() in text_lower:
            return country.name
            
        # Safely check if common_name is part of the text
        if hasattr(country, 'common_name') and country.common_name.lower() in text_lower:
            return country.common_name
    # 2. Prioritize specific regional overrides
    for keyword, country in country_keywords_regional_overrides.items():
        if keyword in text_lower:
            return country
    # 3. Check for broader regions that you want to map to "unknown" or a specific country
    if "north asia" in text_lower or "southeast asia" in text_lower or "east asia" in text_lower:
        return "unknown"

    return "unknown"

# Get the list of English stop words from NLTK
non_meaningful_pop_names = set(stopwords.words('english'))

def parse_population_code_to_country(plain_text_content, table_strings):
    pop_code_country_map = {}
    pop_code_ethnicity_map = {} # NEW: To store ethnicity for structured lookup
    pop_code_specific_loc_map = {} # NEW: To store specific location for structured lookup

    # Regex for parsing population info in structured lists and general text
    # This pattern captures: (Pop Name/Ethnicity) (Pop Code) (Region/Specific Location) (Country) (Linguistic Family)
    # The 'Pop Name/Ethnicity' (Group 1) is often the ethnicity
    pop_info_pattern = re.compile(
          r'([A-Za-z\s]+?)\s+([A-Z]+\d*)\s+'      # Pop Name (Group 1), Pop Code (Group 2) - Changed \d+ to \d* for codes like 'SH'
          r'([A-Za-z\s\(\)\-,\/]+?)\s+'          # Region/Specific Location (Group 3)
          r'(North+|South+|West+|East+|Thailand|Laos|Cambodia|Myanmar|Philippines|Indonesia|Malaysia|China|India|Taiwan|Vietnam|Russia|Nepal|Japan|South Korea)\b' # Country (Group 4)
          r'(?:.*?([A-Za-z\s\-]+))?\s*'          # Optional Linguistic Family (Group 5), made optional with ?, followed by optional space
          r'(\d+(?:\s+\d+\.?\d*)*)?', # Match all the numbers (Group 6) - made optional
          re.IGNORECASE
      )
    for table_str in table_strings:
        table_data = parse_literal_python_list(table_str)
        if table_data:
            is_list_of_lists = bool(table_data) and isinstance(table_data[0], list)
            if is_list_of_lists:
                for row_idx, row in enumerate(table_data):
                    row_text = " ".join(map(str, row))
                    match = pop_info_pattern.search(row_text)
                    if match:
                        pop_name = match.group(1).strip()
                        pop_code = match.group(2).upper()
                        specific_loc_text = match.group(3).strip()
                        country_text = match.group(4).strip()
                        linguistic_family = match.group(5).strip() if match.group(5) else 'unknown'

                        final_country = get_country_from_text(country_text)
                        if final_country == 'unknown': # Try specific loc text for country if direct country is not found
                            final_country = get_country_from_text(specific_loc_text)

                        if pop_code:
                            pop_code_country_map[pop_code] = final_country

                            # Populate ethnicity map (often Pop Name is ethnicity)
                            pop_code_ethnicity_map[pop_code] = pop_name

                            # Populate specific location map
                            pop_code_specific_loc_map[pop_code] = specific_loc_text # Store as is from text
            else:
                row_text = " ".join(map(str, table_data))   
                match = pop_info_pattern.search(row_text)
                if match:
                    pop_name = match.group(1).strip()
                    pop_code = match.group(2).upper()
                    specific_loc_text = match.group(3).strip()
                    country_text = match.group(4).strip()
                    linguistic_family = match.group(5).strip() if match.group(5) else 'unknown'

                    final_country = get_country_from_text(country_text)
                    if final_country == 'unknown': # Try specific loc text for country if direct country is not found
                        final_country = get_country_from_text(specific_loc_text)

                    if pop_code:
                        pop_code_country_map[pop_code] = final_country

                        # Populate ethnicity map (often Pop Name is ethnicity)
                        pop_code_ethnicity_map[pop_code] = pop_name

                        # Populate specific location map
                        pop_code_specific_loc_map[pop_code] = specific_loc_text # Store as is from text

                        # # Special case refinements for ethnicity/location if more specific rules are known from document:
                        # if pop_name.lower() == "khon mueang": # and specific conditions if needed
                        #     pop_code_ethnicity_map[pop_code] = "Khon Mueang"
                        #     # If Khon Mueang has a specific city/district, add here
                        #     # e.g., if 'Chiang Mai' is directly linked to KM1 in a specific table
                        #     # pop_code_specific_loc_map[pop_code] = "Chiang Mai"
                        # elif pop_name.lower() == "lawa":
                        #      pop_code_ethnicity_map[pop_code] = "Lawa"
                        # # Add similar specific rules for other populations (e.g., Mon for MO1, MO2, MO3)
                        # elif pop_name.lower() == "mon":
                        #     pop_code_ethnicity_map[pop_code] = "Mon"
                        #     # For MO2: "West Thailand (Thailand Myanmar border)" -> no city
                        #     # For MO3: "East Myanmar (Thailand Myanmar border)" -> no city
                        #     # If the doc gives "Bangkok" for MO4, add it here for MO4's actual specific_location.
                        # # etc.

    # Fallback to parsing general plain text content (sentences)
    sentences = data_preprocess.extract_sentences(plain_text_content)
    for s in sentences: # Still focusing on just this one sentence
      # Use re.finditer to get all matches
      matches = pop_info_pattern.finditer(s)
      pop_name, pop_code, specific_loc_text, country_text = "unknown", "unknown", "unknown", "unknown"
      for match in matches:
          if match.group(1):
            pop_name = match.group(1).strip()
          if match.group(2):  
            pop_code = match.group(2).upper()
          if match.group(3):  
            specific_loc_text = match.group(3).strip()
          if match.group(4):  
            country_text = match.group(4).strip()
          # linguistic_family = match.group(5).strip() if match.group(5) else 'unknown' # Already captured by pop_info_pattern

          final_country = get_country_from_text(country_text)
          if final_country == 'unknown':
              final_country = get_country_from_text(specific_loc_text)

          if pop_code.lower() not in non_meaningful_pop_names:
            if final_country.lower() not in non_meaningful_pop_names:
              pop_code_country_map[pop_code] = final_country
            if pop_name.lower() not in non_meaningful_pop_names:  
              pop_code_ethnicity_map[pop_code] = pop_name # Default ethnicity from Pop Name
            if specific_loc_text.lower() not in non_meaningful_pop_names:  
              pop_code_specific_loc_map[pop_code] = specific_loc_text

              # Specific rules for ethnicity/location in plain text:
              if pop_name.lower() == "khon mueang":
                  pop_code_ethnicity_map[pop_code] = "Khon Mueang"
              elif pop_name.lower() == "lawa":
                  pop_code_ethnicity_map[pop_code] = "Lawa"
              elif pop_name.lower() == "mon":
                  pop_code_ethnicity_map[pop_code] = "Mon"
              elif pop_name.lower() == "seak": # Added specific rule for Seak
                  pop_code_ethnicity_map[pop_code] = "Seak"
              elif pop_name.lower() == "nyaw": # Added specific rule for Nyaw
                  pop_code_ethnicity_map[pop_code] = "Nyaw"
              elif pop_name.lower() == "nyahkur": # Added specific rule for Nyahkur
                  pop_code_ethnicity_map[pop_code] = "Nyahkur"
              elif pop_name.lower() == "suay": # Added specific rule for Suay
                  pop_code_ethnicity_map[pop_code] = "Suay"
              elif pop_name.lower() == "soa": # Added specific rule for Soa
                  pop_code_ethnicity_map[pop_code] = "Soa"
              elif pop_name.lower() == "bru": # Added specific rule for Bru
                  pop_code_ethnicity_map[pop_code] = "Bru"
              elif pop_name.lower() == "khamu": # Added specific rule for Khamu
                  pop_code_ethnicity_map[pop_code] = "Khamu"

    return pop_code_country_map, pop_code_ethnicity_map, pop_code_specific_loc_map

def general_parse_population_code_to_country(plain_text_content, table_strings):
    pop_code_country_map = {}
    pop_code_ethnicity_map = {}
    pop_code_specific_loc_map = {}
    sample_id_to_pop_code = {}

    for table_str in table_strings:
        table_data = parse_literal_python_list(table_str)
        if not table_data or not isinstance(table_data[0], list):
            continue

        header_row = [col.lower() for col in table_data[0]]
        header_map = {col: idx for idx, col in enumerate(header_row)}

        # MJ17: Direct PopCode → Country
        if 'id' in header_map and 'country' in header_map:
            for row in table_strings[1:]:
                row = parse_literal_python_list(row)[0]
                if len(row) < len(header_row):
                    continue
                pop_code = str(row[header_map['id']]).strip()
                country = str(row[header_map['country']]).strip()
                province = row[header_map['province']].strip() if 'province' in header_map else 'unknown'
                pop_group = row[header_map['population group / region']].strip() if 'population group / region' in header_map else 'unknown'
                pop_code_country_map[pop_code] = country
                pop_code_specific_loc_map[pop_code] = province
                pop_code_ethnicity_map[pop_code] = pop_group

        # A1YU101 or EBK/KSK: SampleID → PopCode
        elif 'sample id' in header_map and 'population code' in header_map:
            for row in table_strings[1:]:
                row = parse_literal_python_list(row)[0]
                if len(row) < 2:
                    continue
                sample_id = row[header_map['sample id']].strip().upper()
                pop_code = row[header_map['population code']].strip().upper()
                sample_id_to_pop_code[sample_id] = pop_code

        # PopCode → Country (A1YU101/EBK mapping)
        elif 'population code' in header_map and 'country' in header_map:
            for row in table_strings[1:]:
                row = parse_literal_python_list(row)[0]
                if len(row) < 2:
                    continue
                pop_code = row[header_map['population code']].strip().upper()
                country = row[header_map['country']].strip()
                pop_code_country_map[pop_code] = country

    return pop_code_country_map, pop_code_ethnicity_map, pop_code_specific_loc_map, sample_id_to_pop_code

def chunk_text(text, chunk_size=500, overlap=50):
    """Splits text into chunks (by words) with overlap."""
    chunks = []
    words = text.split()
    num_words = len(words)

    start = 0
    while start < num_words:
        end = min(start + chunk_size, num_words)
        chunk = " ".join(words[start:end])
        chunks.append(chunk)

        if end == num_words:
            break
        start += chunk_size - overlap # Move start by (chunk_size - overlap)
    return chunks

def build_vector_index_and_data(doc_path, index_path="faiss_index.bin", chunks_path="document_chunks.json", structured_path="structured_lookup.json"):
    """
    Reads document, builds structured lookup, chunks remaining text, embeds chunks,
    and builds/saves a FAISS index.
    """
    print("Step 1: Reading document and extracting structured data...")
    # plain_text_content, table_strings, document_title = read_docx_text(doc_path) # Get document_title here

    # sample_id_map, contiguous_ranges_data = parse_sample_id_to_population_code(plain_text_content)
    # pop_code_to_country, pop_code_to_ethnicity, pop_code_to_specific_loc = parse_population_code_to_country(plain_text_content, table_strings)

    # master_structured_lookup = {}
    # master_structured_lookup['document_title'] = document_title # Store document title
    # master_structured_lookup['sample_id_map'] = sample_id_map
    # master_structured_lookup['contiguous_ranges'] = dict(contiguous_ranges_data)
    # master_structured_lookup['pop_code_to_country'] = pop_code_to_country
    # master_structured_lookup['pop_code_to_ethnicity'] = pop_code_to_ethnicity # NEW: Store pop_code to ethnicity map
    # master_structured_lookup['pop_code_to_specific_loc'] = pop_code_to_specific_loc # NEW: Store pop_code to specific_loc map


    # # Final consolidation: Use sample_id_map to derive full info for queries
    # final_structured_entries = {}
    # for sample_id, pop_code in master_structured_lookup['sample_id_map'].items():
    #     country = master_structured_lookup['pop_code_to_country'].get(pop_code, 'unknown')
    #     ethnicity = master_structured_lookup['pop_code_to_ethnicity'].get(pop_code, 'unknown') # Retrieve ethnicity
    #     specific_location = master_structured_lookup['pop_code_to_specific_loc'].get(pop_code, 'unknown') # Retrieve specific location

    #     final_structured_entries[sample_id] = {
    #         'population_code': pop_code,
    #         'country': country,
    #         'type': 'modern',
    #         'ethnicity': ethnicity, # Store ethnicity
    #         'specific_location': specific_location # Store specific location
    #     }
    # master_structured_lookup['final_structured_entries'] = final_structured_entries
    plain_text_content, table_strings, document_title = read_docx_text(doc_path)
    pop_code_to_country, pop_code_to_ethnicity, pop_code_to_specific_loc, sample_id_map = general_parse_population_code_to_country(plain_text_content, table_strings)

    final_structured_entries = {}
    if sample_id_map:
        for sample_id, pop_code in sample_id_map.items():
            country = pop_code_to_country.get(pop_code, 'unknown')
            ethnicity = pop_code_to_ethnicity.get(pop_code, 'unknown')
            specific_loc = pop_code_to_specific_loc.get(pop_code, 'unknown')
            final_structured_entries[sample_id] = {
                'population_code': pop_code,
                'country': country,
                'type': 'modern',
                'ethnicity': ethnicity,
                'specific_location': specific_loc
            }
    else:
        for pop_code in pop_code_to_country.keys():
            country = pop_code_to_country.get(pop_code, 'unknown')
            ethnicity = pop_code_to_ethnicity.get(pop_code, 'unknown')
            specific_loc = pop_code_to_specific_loc.get(pop_code, 'unknown')
            final_structured_entries[pop_code] = {
                'population_code': pop_code,
                'country': country,
                'type': 'modern',
                'ethnicity': ethnicity,
                'specific_location': specific_loc
            }
    if not final_structured_entries:
      # traditional way of A1YU101
      sample_id_map, contiguous_ranges_data = parse_sample_id_to_population_code(plain_text_content)
      pop_code_to_country, pop_code_to_ethnicity, pop_code_to_specific_loc = parse_population_code_to_country(plain_text_content, table_strings)
      if sample_id_map:
        for sample_id, pop_code in sample_id_map.items():
            country = pop_code_to_country.get(pop_code, 'unknown')
            ethnicity = pop_code_to_ethnicity.get(pop_code, 'unknown')
            specific_loc = pop_code_to_specific_loc.get(pop_code, 'unknown')
            final_structured_entries[sample_id] = {
                'population_code': pop_code,
                'country': country,
                'type': 'modern',
                'ethnicity': ethnicity,
                'specific_location': specific_loc
            }
      else:
          for pop_code in pop_code_to_country.keys():
              country = pop_code_to_country.get(pop_code, 'unknown')
              ethnicity = pop_code_to_ethnicity.get(pop_code, 'unknown')
              specific_loc = pop_code_to_specific_loc.get(pop_code, 'unknown')
              final_structured_entries[pop_code] = {
                  'population_code': pop_code,
                  'country': country,
                  'type': 'modern',
                  'ethnicity': ethnicity,
                  'specific_location': specific_loc
              }
    
    master_lookup = {
        'document_title': document_title,
        'pop_code_to_country': pop_code_to_country,
        'pop_code_to_ethnicity': pop_code_to_ethnicity,
        'pop_code_to_specific_loc': pop_code_to_specific_loc,
        'sample_id_map': sample_id_map,
        'final_structured_entries': final_structured_entries
    }
    print(f"Structured lookup built with {len(final_structured_entries)} entries in 'final_structured_entries'.")

    with open(structured_path, 'w') as f:
        json.dump(master_lookup, f, indent=4)
    print(f"Structured lookup saved to {structured_path}.")

    print("Step 2: Chunking document for RAG vector index...")
    # replace the chunk here with the all_output from process_inputToken and fallback to this traditional chunk
    clean_text, clean_table = "", ""
    if plain_text_content:
      clean_text = data_preprocess.normalize_for_overlap(plain_text_content)
    if table_strings:
      clean_table = data_preprocess.normalize_for_overlap(". ".join(table_strings))
    all_clean_chunk = clean_text + clean_table
    document_chunks = chunk_text(all_clean_chunk)
    print(f"Document chunked into {len(document_chunks)} chunks.")
    
    print("Step 3: Generating embeddings for chunks (this might take time and cost API calls)...")

    embedding_model_for_chunks = genai.GenerativeModel('models/text-embedding-004')

    chunk_embeddings = []
    for i, chunk in enumerate(document_chunks):
        embedding = get_embedding(chunk, task_type="RETRIEVAL_DOCUMENT")
        if embedding is not None and embedding.shape[0] > 0:
            chunk_embeddings.append(embedding)
        else:
            print(f"Warning: Failed to get valid embedding for chunk {i}. Skipping.")
            chunk_embeddings.append(np.zeros(768, dtype='float32'))

    if not chunk_embeddings:
        raise ValueError("No valid embeddings generated. Check get_embedding function and API.")

    embedding_dimension = chunk_embeddings[0].shape[0]
    index = faiss.IndexFlatL2(embedding_dimension)
    index.add(np.array(chunk_embeddings))

    faiss.write_index(index, index_path)
    with open(chunks_path, "w") as f:
        json.dump(document_chunks, f)

    print(f"FAISS index built and saved to {index_path}.")
    print(f"Document chunks saved to {chunks_path}.")
    return master_lookup, index, document_chunks, all_clean_chunk


def load_rag_assets(index_path="faiss_index.bin", chunks_path="document_chunks.json", structured_path="structured_lookup.json"):
    """Loads pre-built RAG assets (FAISS index, chunks, structured lookup)."""
    print("Loading RAG assets...")
    master_structured_lookup = {}
    if os.path.exists(structured_path):
        with open(structured_path, 'r') as f:
            master_structured_lookup = json.load(f)
        print("Structured lookup loaded.")
    else:
        print("Structured lookup file not found. Rebuilding is likely needed.")

    index = None
    chunks = []
    if os.path.exists(index_path) and os.path.exists(chunks_path):
        try:
            index = faiss.read_index(index_path)
            with open(chunks_path, "r") as f:
                chunks = json.load(f)
            print("FAISS index and chunks loaded.")
        except Exception as e:
            print(f"Error loading FAISS index or chunks: {e}. Will rebuild.")
            index = None
            chunks = []
    else:
        print("FAISS index or chunks files not found.")

    return master_structured_lookup, index, chunks
# Helper function for query_document_info
def exactInContext(text, keyword):
# try keyword_prfix
  # code_pattern = re.compile(r'([A-Z0-9]+?)(\d+)$', re.IGNORECASE)
  # # Attempt to parse the keyword into its prefix and numerical part using re.search
  # keyword_match = code_pattern.search(keyword)
  # keyword_prefix = None
  # keyword_num = None
  # if keyword_match:
  #     keyword_prefix = keyword_match.group(1).lower()
  #     keyword_num = int(keyword_match.group(2))
  text = text.lower()
  idx = text.find(keyword.lower())
  if idx == -1:
    # if keyword_prefix:
    #   idx = text.find(keyword_prefix)
    # if idx == -1:
    #   return False
    return False
  return True
def chooseContextLLM(contexts, kw):
  # if kw in context
  for con in contexts:
    context = contexts[con]
    if context:
      if exactInContext(context, kw):
        return con, context    
  #if cannot find anything related to kw in context, return all output
  if contexts["all_output"]:
    return "all_output", contexts["all_output"]
  else:
    # if all_output not exist
    # look of chunk and still not exist return document chunk
    if contexts["chunk"]: return "chunk", contexts["chunk"]
    elif contexts["document_chunk"]:  return "document_chunk", contexts["document_chunk"]
    else: return None, None  
def clean_llm_output(llm_response_text, output_format_str):
    results = []
    lines = llm_response_text.strip().split('\n')
    output_country, output_type, output_ethnicity, output_specific_location = [],[],[],[]
    for line in lines:
        extracted_country, extracted_type, extracted_ethnicity, extracted_specific_location = "unknown", "unknown", "unknown", "unknown"
        line = line.strip()
        if output_format_str == "ethnicity, specific_location/unknown": # Targeted RAG output
            parsed_output = re.search(r'^\s*([^,]+?),\s*(.+?)\s*$', llm_response_text)
            if parsed_output:
                extracted_ethnicity = parsed_output.group(1).strip()
                extracted_specific_location = parsed_output.group(2).strip()
            else:
                print("  DEBUG: LLM did not follow expected 2-field format for targeted RAG. Defaulting to unknown for ethnicity/specific_location.")
                extracted_ethnicity = 'unknown'
                extracted_specific_location = 'unknown'
        elif output_format_str == "modern/ancient/unknown, ethnicity, specific_location/unknown":
          parsed_output = re.search(r'^\s*([^,]+?),\s*([^,]+?),\s*(.+?)\s*$', llm_response_text)
          if parsed_output:
              extracted_type = parsed_output.group(1).strip()
              extracted_ethnicity = parsed_output.group(2).strip()
              extracted_specific_location = parsed_output.group(3).strip()
          else:
              # Fallback: check if only 2 fields
              parsed_output_2_fields = re.search(r'^\s*([^,]+?),\s*([^,]+?)\s*$', llm_response_text)
              if parsed_output_2_fields:
                  extracted_type = parsed_output_2_fields.group(1).strip()
                  extracted_ethnicity = parsed_output_2_fields.group(2).strip()
                  extracted_specific_location = 'unknown'
              else:
                  # even simpler fallback: 1 field only
                  parsed_output_1_field = re.search(r'^\s*([^,]+?)\s*$', llm_response_text)
                  if parsed_output_1_field:
                      extracted_type = parsed_output_1_field.group(1).strip()
                      extracted_ethnicity = 'unknown'
                      extracted_specific_location = 'unknown'
                  else:
                      print("  DEBUG: LLM did not follow any expected simplified format. Attempting verbose parsing fallback.")
                      type_match_fallback = re.search(r'Type:\s*([A-Za-z\s-]+)', llm_response_text)
                      extracted_type = type_match_fallback.group(1).strip() if type_match_fallback else 'unknown'
                      extracted_ethnicity = 'unknown'
                      extracted_specific_location = 'unknown'
        else:
          parsed_output = re.search(r'^\s*([^,]+?),\s*([^,]+?),\s*([^,]+?),\s*(.+?)\s*$', line)
          if parsed_output:
              extracted_country = parsed_output.group(1).strip()
              extracted_type = parsed_output.group(2).strip()
              extracted_ethnicity = parsed_output.group(3).strip()
              extracted_specific_location = parsed_output.group(4).strip()
          else:
              print(f"  DEBUG: Line did not follow expected 4-field format: {line}")
              parsed_output_2_fields = re.search(r'^\s*([^,]+?),\s*([^,]+?)\s*$', line)
              if parsed_output_2_fields:
                  extracted_country = parsed_output_2_fields.group(1).strip()
                  extracted_type = parsed_output_2_fields.group(2).strip()
                  extracted_ethnicity = 'unknown'
                  extracted_specific_location = 'unknown'
              else:
                  print(f"  DEBUG: Fallback to verbose-style parsing: {line}")
                  country_match_fallback = re.search(r'Country:\s*([A-Za-z\s-]+)', line)
                  type_match_fallback = re.search(r'Type:\s*([A-Za-z\s-]+)', line)
                  extracted_country = country_match_fallback.group(1).strip() if country_match_fallback else 'unknown'
                  extracted_type = type_match_fallback.group(1).strip() if type_match_fallback else 'unknown'
                  extracted_ethnicity = 'unknown'
                  extracted_specific_location = 'unknown'

        results.append({
            "country": extracted_country,
            "type": extracted_type,
            "ethnicity": extracted_ethnicity,
            "specific_location": extracted_specific_location
            #"country_explain":extracted_country_explain,
            #"type_explain": extracted_type_explain
        })
    # if more than 2 results
    if output_format_str == "ethnicity, specific_location/unknown":
      for result in results:
        if result["ethnicity"] not in output_ethnicity:
          output_ethnicity.append(result["ethnicity"])
        if result["specific_location"] not in output_specific_location:  
          output_specific_location.append(result["specific_location"])
      return " or ".join(output_ethnicity), " or ".join(output_specific_location)     
    elif output_format_str == "modern/ancient/unknown, ethnicity, specific_location/unknown":
      for result in results:
        if result["type"] not in output_type:
          output_type.append(result["type"])
        if result["ethnicity"] not in output_ethnicity:
          output_ethnicity.append(result["ethnicity"])
        if result["specific_location"] not in output_specific_location:  
          output_specific_location.append(result["specific_location"])

      return " or ".join(output_type)," or ".join(output_ethnicity), " or ".join(output_specific_location)    
    else:
      for result in results:
        if result["country"] not in output_country:
          output_country.append(result["country"])
        if result["type"] not in output_type:
          output_type.append(result["type"])
        if result["ethnicity"] not in output_ethnicity:
          output_ethnicity.append(result["ethnicity"])
        if result["specific_location"] not in output_specific_location:  
          output_specific_location.append(result["specific_location"])
      return " or ".join(output_country)," or ".join(output_type)," or ".join(output_ethnicity), " or ".join(output_specific_location)           

# def parse_multi_sample_llm_output(raw_response: str, output_format_str):
#     """
#     Parse LLM output with possibly multiple metadata lines + shared explanations.
#     """
#     lines = [line.strip() for line in raw_response.strip().splitlines() if line.strip()]
#     metadata_list = []
#     explanation_lines = []
#     if output_format_str == "country_name, modern/ancient/unknown":
#         parts = [x.strip() for x in lines[0].split(",")]
#         if len(parts)==2:
#           metadata_list.append({
#               "country": parts[0],
#               "sample_type": parts[1]#,
#               #"ethnicity": parts[2],
#               #"location": parts[3]
#           })
#         if 1<len(lines):
#           line = lines[1]
#           if "\n" in line:  line = line.split("\n")
#           if ". " in line: line = line.split(". ")
#           if isinstance(line,str): line = [line]
#           explanation_lines += line
#     elif output_format_str == "modern/ancient/unknown":
#       metadata_list.append({
#           "country": "unknown",
#           "sample_type": lines[0]#,
#           #"ethnicity": parts[2],
#           #"location": parts[3]
#       })
#       explanation_lines.append(lines[1])

#     # Assign explanations (optional) to each sample — same explanation reused
#     for md in metadata_list:
#         md["country_explanation"] = None
#         md["sample_type_explanation"] = None

#         if md["country"].lower() != "unknown" and len(explanation_lines) >= 1:
#             md["country_explanation"] = explanation_lines[0]

#         if md["sample_type"].lower() != "unknown":
#             if len(explanation_lines) >= 2:
#                 md["sample_type_explanation"] = explanation_lines[1]
#             elif len(explanation_lines) == 1 and md["country"].lower() == "unknown":
#                 md["sample_type_explanation"] = explanation_lines[0]
#             elif len(explanation_lines) == 1:
#                 md["sample_type_explanation"] = explanation_lines[0]
#     return metadata_list

def parse_multi_sample_llm_output(raw_response: str, output_format_str):
    """
    Parse LLM output for one sample's fields.

    Primary format (what multi_prompts() now asks for): one block per field,
    reasoning written BEFORE the terse answer, so the model can't commit a
    value before it has reasoned about it --

        FIELD: <field_name>
        REASONING: <narrative> [Sources: ...] [Conflict: ...] [ID-match: true|false]
        ANSWER: <value>

    Falls back to the older "Line 1 pipe-separated summary + per-field
    explanation lines" layout (kept verbatim below) if the model doesn't
    fully comply with the block format -- robustness net, not the norm.
    """
    output_formats = output_format_str.split(", ") if output_format_str else []

    # ── Primary: FIELD: / REASONING: / ANSWER: blocks ──────────────────────────
    block_pattern = re.compile(
        r'FIELD:\s*(.+?)\s*\n'
        r'REASONING:\s*(.+?)\s*\n'
        r'ANSWER:\s*(.+?)\s*(?=\n\s*FIELD:|\Z)',
        re.DOTALL,
    )
    blocks = block_pattern.findall(raw_response)

    if blocks:
        field_map = {}
        for field_name, reasoning, answer in blocks:
            field_map[field_name.strip().lower()] = (reasoning.strip(), answer.strip())

        metadata_list = {}
        for output in output_formats:
            metadata_list[output] = {"answer": "", output + "_explanation": ""}
            match = field_map.get(output.lower())
            if match is None:
                metadata_list[output]["answer"] = "unknown"
                metadata_list[output][output + "_explanation"] = "unknown"
                continue
            reasoning, answer = match
            # Tolerate the model echoing "field_name: value" inside ANSWER despite instructions
            if ": " in answer and answer.split(": ", 1)[0].strip().lower() == output.lower():
                answer = answer.split(": ", 1)[1].strip()
            if not answer or "unknown" in answer.lower():
                metadata_list[output]["answer"] = "unknown"
                metadata_list[output][output + "_explanation"] = "unknown"
            else:
                metadata_list[output]["answer"] = answer
                metadata_list[output][output + "_explanation"] = reasoning

        print("parsed metadata_list keys (block format):", list(metadata_list.keys()))
        return metadata_list

    # ── Fallback: older "Line 1 summary + explanation lines" layout ────────────
    metadata_list = {}
    raw_lines = raw_response.strip().split("\n")
    first_line = raw_lines[0].strip() if raw_lines else ""
    explanation_lines_raw = [x for x in raw_lines[1:] if x.strip()]

    # Use pipe ' | ' as primary separator (avoids splitting comma-containing CONFLICT values)
    if ' | ' in first_line:
        output_answers = [x.strip() for x in first_line.split(' | ')]
    else:
        output_answers = re.split(r",\s*", first_line)

    # ── Build per-field explanation map ───────────────────────────────────────
    # Strategy A: try **field_name:** markers anywhere in the explanation block
    full_expl_text = " ".join(explanation_lines_raw)
    field_expl_map: dict = {}
    for fmt in output_formats:
        escaped = re.escape(fmt)
        # Match **field:** ... up to next **field:** or end
        pattern = rf'\*{{1,2}}{escaped}\*{{0,2}}\s*[:\-]?\s*(.+?)(?=\*{{1,2}}[A-Za-z_/]+\*{{0,2}}\s*[:\-]|$)'
        m = re.search(pattern, full_expl_text, re.IGNORECASE | re.DOTALL)
        if m:
            sentence = m.group(1).strip().split("\n")[0]  # first sentence only
            field_expl_map[fmt] = sentence

    # Strategy B: ordered lines (one per field)
    ordered_lines = explanation_lines_raw
    if not field_expl_map and len(ordered_lines) == 1 and ". " in ordered_lines[0]:
        line = ordered_lines[0]
        # Don't split when rich citation tags are present — they may contain '. '
        if '[Sources:' not in line and '[Conflict:' not in line and '[Source:' not in line:
            ordered_lines = [s.strip() for s in line.split(". ") if s.strip()]

    # ── Assign answers + per-field explanations ───────────────────────────────
    for o, output in enumerate(output_formats):
        metadata_list[output] = {"answer": "", output + "_explanation": ""}

        # Answer
        if o < len(output_answers):
            ans = output_answers[o].strip()
            try:
                if ": " in ans:
                    ans = ans.split(": ", 1)[1]
            except Exception:
                pass
            metadata_list[output]["answer"] = ans
            if "unknown" in metadata_list[output]["answer"].lower():
                metadata_list[output]["answer"] = "unknown"
        else:
            metadata_list[output]["answer"] = "unknown"

        # Explanation — one sentence, assigned to this field specifically
        if metadata_list[output]["answer"] != "unknown":
            if output in field_expl_map:
                explain = field_expl_map[output]
            elif o < len(ordered_lines):
                explain = ordered_lines[o]
            elif ordered_lines:
                explain = ordered_lines[-1]
            else:
                explain = ""
            # Strip leading **field:** prefix if present
            explain = re.sub(r'^\*{1,2}[A-Za-z_/\-]+\*{0,2}\s*[:\-]\s*', '', explain).strip()
            metadata_list[output][output + "_explanation"] = explain
        else:
            metadata_list[output][output + "_explanation"] = "unknown"

    print("parsed metadata_list keys (fallback format):", list(metadata_list.keys()))
    return metadata_list


_PROMPT_HEADER_RE = re.compile(
    r'===\s*ANSWERS FOR PROMPT\s+(\d+)\s*\(accession\s+(.+?)\)\s*===',
    re.IGNORECASE,
)


def split_batched_llm_response(raw_response: str, accs: list) -> dict:
    """
    Split one LLM response covering multiple accessions (multi_prompts()'s
    '=== ANSWERS FOR PROMPT N (accession X) ===' headers) into per-accession
    raw-text segments.

    Primary mapping: the header's own captured prompt number N -> accs[N-1]
    (multi_prompts() builds "Prompt {acc_pos+1}" for accs[acc_pos], so N-1 is
    that accession's position in the original request). This holds regardless
    of what order the model answers the prompts in -- unlike matching by
    header position in the text, which would silently misattribute data if
    the model ever reordered its answers while still labeling headers
    correctly.

    Falls back to positional order (i-th header found -> accs[i]) only when a
    header's number is missing, unparseable, or out of range for this batch --
    e.g. the model garbles the number but still separates answers in order.
    A header whose echoed accession text (group 2) disagrees with the
    resolved index is logged but still used -- the number is the source of
    truth, the echoed text is only a human-readable label the model might
    paraphrase.

    Falls back to returning the whole response under accs[0] when no header
    is found at all -- preserves today's single-accession behavior (and the
    older pipe-separated fallback format parse_multi_sample_llm_output
    already handles) for the batch-of-1 case.
    """
    matches = list(_PROMPT_HEADER_RE.finditer(raw_response))
    if not matches:
        return {accs[0]: raw_response} if accs else {}

    segments = {}
    for i, m in enumerate(matches):
        seg_start = m.end()
        seg_end = matches[i + 1].start() if i + 1 < len(matches) else len(raw_response)

        acc = None
        try:
            n = int(m.group(1))
            if 1 <= n <= len(accs):
                acc = accs[n - 1]
        except (TypeError, ValueError):
            pass

        if acc is None:
            if i < len(accs):
                acc = accs[i]
                print(f"[split_batched_llm_response] header #{i+1} has missing/invalid/"
                      f"out-of-range prompt number ({m.group(1)!r}) -- falling back to "
                      f"positional match: {acc}")
            else:
                print(f"[split_batched_llm_response] header #{i+1} unmatchable "
                      f"(number {m.group(1)!r} invalid, no positional slot left) -- dropped")
                continue
        elif acc != m.group(2).strip() and acc.split('.')[0] != m.group(2).strip().split('.')[0]:
            print(f"[split_batched_llm_response] header says accession "
                  f"{m.group(2)!r} but prompt number {m.group(1)} maps to {acc!r} "
                  f"-- using the number (source of truth)")

        if acc in segments:
            print(f"[split_batched_llm_response] header #{i+1} resolved to {acc!r}, "
                  f"which already has a segment from an earlier header -- overwriting "
                  f"(duplicate/colliding prompt number; last write wins)")
        segments[acc] = raw_response[seg_start:seg_end]
    return segments


def merge_metadata_outputs(metadata_list):
    """
    Merge a list of metadata dicts into one, combining differing values with 'or'.
    Assumes all dicts have the same keys.
    """
    if not metadata_list:
        return {}

    merged = {}
    keys = metadata_list[0].keys()

    for key in keys:
        values = [md[key] for md in metadata_list if key in md]
        unique_values = list(dict.fromkeys(values))  # preserve order, remove dupes
        if "unknown" in unique_values:
          unique_values.pop(unique_values.index("unknown"))
        if len(unique_values) == 1:
            merged[key] = unique_values[0]
        else:
            merged[key] = " or ".join(unique_values)

    return merged

import time
import random

def safe_call_llm(prompt, model="gemini-2.5-flash-lite", max_retries=5):
    retry_delay = 20
    for attempt in range(max_retries):
        try:
            resp_text, resp_model = call_llm_api(prompt, model)
            return resp_text, resp_model
        except Exception as e:
            error_msg = str(e)
            if "429" in error_msg or "quota" in error_msg.lower() or "rate" in error_msg.lower() or "overloaded" in error_msg.lower():
                print(f"\n⚠️ Rate limit hit (attempt {attempt+1}/{max_retries}).")

                retry_after = None
                for word in error_msg.split():
                    if "retry" in word.lower() and "s" in word:
                        try:
                            retry_after = float(word.replace("s","").replace(".",""))
                        except:
                            pass

                wait_time = retry_after if retry_after else retry_delay
                print(f"⏳ Waiting {wait_time:.1f} seconds before retrying...")
                time.sleep(wait_time)

                retry_delay *= 2
            else:
                raise e

    raise RuntimeError("❌ Failed after max retries because of repeated rate limits.")

def outputs_from_multiPrompts(raw_response: str, output_format_str, acc_prompts):
  # Split the text based on the pattern '**Prompt X:'
  raw_response = re.split(r'\*\*Prompt \d+:', text)

  # Remove any empty sections from the split list
  prompts = [prompt.strip() for prompt in raw_response if prompt.strip()]

  # Create a list of output strings
  outputs = {}
  accs = list(acc_prompts.keys())
  # Loop through the prompts and combine the header and body
  for i in range(0, len(prompts)):
      prompt_header = prompts[i].strip()  # This is the "USA, unknown, Venezuela" or similar part
      prompt_header = re.sub(r'^\*\*\n', '', prompt_header)  # Remove any leading '**\n'
      accession, output = accs[i], ""
      if i + 1 < len(prompts):  # Check if there is a next body text
          prompt_body = prompts[i + 1].strip()  # This is the body of the response
          # Remove any unwanted '**\n' before the prompt content
          output = f"{prompt_header}\n\n{prompt_body}"
      else:
          # If no body exists, add only the header (though this case shouldn't occur in this example)
          output = f"{prompt_header}\n\n"
      metadata_list = parse_multi_sample_llm_output(output, output_format_str)
      outputs[accession] = metadata_list    
  return outputs   


# ── Built-in default standardization schema ────────────────────────────────────
# Used by multi_prompts() whenever the caller doesn't supply its own
# standardization_schema dict (e.g. no schema URL was given for this run) --
# gives every run a sensible, general baseline of field definitions/allowed
# values instead of no schema guidance at all. Sourced from curatedMetagenomicData's
# own data dictionary (github.com/waldronlab/curatedMetagenomicDataCuration).
_DEFAULT_SCHEMA_CSV_TEXT = r""""col.name","col.class","unique","required","multiplevalues","description","allowedvalues","static.enum","dynamic.enum","dynamic.enum.property","delimiter","separator","corpus.type","display.order","display.group"
"study_name","character","non-unique","required",FALSE,"Canonical study identifier (for example EinsteinA_YYYY).","[a-zA-Z-]+_[0-9]{4}|[a-zA-Z-]+_[0-9]{4}[a-zA-Z-]+|[a-zA-Z-]+_[0-9]{4}_[a-zA-Z-]+|[a-zA-Z-]+_[0-9]{4}_[a-zA-Z0-9]+",NA,NA,NA,NA,NA,"regexp",1,"00 Identifiers"
"subject_id","character","non-unique","required",FALSE,"Unique identifier for a participant within a study. Values may be non-unique e.g. when there are repeated measures for a participant.","[0-9a-zA-Z]\S+",NA,NA,NA,NA,NA,"regexp",2,"00 Identifiers"
"sample_id","character","unique","required",FALSE,"Unique identifier for a biospecimen/sample.","[0-9a-zA-Z]\S+",NA,NA,NA,NA,NA,"regexp",3,"00 Identifiers"
"curator","character","non-unique","optional",TRUE,"Name(s) of curator(s) responsible for this entry.",NA,NA,NA,NA,";",NA,"any",4,"00 Identifiers"
"pmid","integer","non-unique","optional",FALSE,"PubMed identifier of the primary publication for the study.","[0-9]{8}",NA,NA,NA,NA,NA,"integer",5,"00 Identifiers"
"subcohort","character","non-unique","optional",FALSE,"Identifier for a subset of the study cohort.","[0-9a-zA-Z]\S+",NA,NA,NA,NA,NA,"regexp",6,"00 Identifiers"
"target_condition","character","non-unique","required",TRUE,"Primary phenotype(s) or condition(s) investigated in the study. Lookup seed ontology IDs: NCIT:C7057; MONDO:0000001. Type: looked_up. Looked up: Yes. Includes child terms: Yes. Existing-dataset-only values now: Conditional (Yes only when lookup falls back to curated existing dataset values)",NA,"NCIT:C115935","NCIT:C7057;MONDO:0000001","descendant",";",NA,"dynamic_enum;static_enum",1,"01 Sampling / Study"
"control","character","non-unique","required",FALSE,"Case-control role of the sample in the study. Static ontology IDs: NCIT:C142703; NCIT:C49152; NCIT:C69062. Type: static.","Study Control|Case|Not Used","NCIT:C142703|NCIT:C49152|NCIT:C69062",NA,NA,NA,NA,"static_enum",2,"01 Sampling / Study"
"country","character","non-unique","optional",FALSE,"Country where the participant resides and/or the sample was collected. Lookup seed ontology IDs: NCIT:C25464. Type: looked_up. Looked up: Yes. Includes child terms: Yes. Existing-dataset-only values now: Conditional (Yes only when lookup falls back to curated existing dataset values)",NA,NA,"NCIT:C25464","descendant",NA,NA,"dynamic_enum",3,"01 Sampling / Study"
"body_site","character","non-unique","required",FALSE,"Anatomical body site from which the sample was collected. Static ontology IDs: UBERON:0001988; UBERON:0000167; UBERON:0001003; UBERON:0000996; UBERON:0001707; UBERON:0001913. Type: static.","feces|oral cavity|skin epidermis|vagina|nasal cavity|milk","UBERON:0001988|UBERON:0000167|UBERON:0001003|UBERON:0000996|UBERON:0001707|UBERON:0001913",NA,NA,NA,NA,"static_enum",4,"01 Sampling / Study"
"body_site_details","character","non-unique","optional",TRUE,"More specific anatomical subsite for sample collection than body_site. Lookup seed ontology IDs: UBERON:0001062. Type: looked_up. Looked up: Yes. Includes child terms: Yes. Existing-dataset-only values now: Conditional (Yes only when lookup falls back to curated existing dataset values)",NA,NA,"UBERON:0001062",NA,";",NA,"dynamic_enum",5,"01 Sampling / Study"
"days_from_first_collection","integer","non-unique","optional",FALSE,"Days elapsed since the first collection time point in longitudinal studies.","[0-9]+",NA,NA,NA,NA,NA,"integer",6,"01 Sampling / Study"
"lifestyle","character","non-unique","optional",FALSE,"Lifestyle category (primarily used for traditional or non-westernized populations).","Hunter-gatherer|Agriculturalist|Agropastoralist|Pastoralist|Fisher",NA,NA,NA,NA,NA,"custom_enum",7,"01 Sampling / Study"
"location","character","non-unique","optional",FALSE,"Free-text additional location detail (for example city or region).",".+",NA,NA,NA,NA,NA,"any",8,"01 Sampling / Study"
"probing_pocket_depth","character","non-unique","optional",TRUE,"Tooth surface/region where periodontal pocket depth was measured. Static ontology IDs: FMA:64849; FMA:55649; FMA:55647; FMA:55650. Type: static.","Buccal surface|Distal surface of tooth|Lingual surface of tooth|Mesial surface of tooth","FMA:64849|FMA:55649|FMA:55647|FMA:55650",NA,NA,"<;>",NA,"static_enum",9,"01 Sampling / Study"
"westernized","character","non-unique","optional",FALSE,"Whether the participant is classified as westernized.","Yes|No",NA,NA,NA,NA,NA,"binary",10,"01 Sampling / Study"
"antibiotics_exclusion_period","double","non-unique","optional",FALSE,"Minimum washout period from antibiotic use required by study criteria.","[0-9]+\.?[0-9]*",NA,NA,NA,NA,NA,"regexp",11,"01 Sampling / Study"
"antibiotics_exclusion_period_unit","character","non-unique","optional",FALSE,"Unit used for `antibiotics_exclusion_period`. Static ontology IDs: NCIT:C25301; NCIT:C29844; NCIT:C29846. Type: static.","Day|Week|Month","NCIT:C25301|NCIT:C29844|NCIT:C29846",NA,NA,NA,NA,"static_enum",12,"01 Sampling / Study"
"probing_pocket_depth_value","character","non-unique","optional",TRUE,"Numeric probing depth measurement corresponding to `probing_pocket_depth`.","^\d+(\.\d+)?$",NA,NA,NA,"<;>",NA,"character",13,"01 Sampling / Study"
"dna_extraction_kit","character","non-unique","optional",FALSE,"Name of the DNA extraction kit or protocol used.","Qiagen|Gnome|MoBio|MPBio|NorgenBiotek|Illuminakit|Maxwell_LEV|PSP_Spin_Stool|Tiangen|PowerSoil|Chemagen|other|PowerSoilPro|ZR_Fecal_DNA_MiniPrep|KAMA_Hyper_Prep|thermo_fisher|QIAamp",NA,NA,NA,NA,NA,"custom_enum",1,"02 Sequencing / QC"
"ncbi_accession","character","non-unique","optional",FALSE,"NCBI SRA/ENA accession(s) associated with the sample or run.","[ES]R[SR][0-9]+",NA,NA,NA,NA,NA,"regexp",4,"02 Sequencing / QC"
"sequencing_platform","character","non-unique","required",FALSE,"Sequencing platform/instrument used to generate the data.","IlluminaHiSeq|IlluminaMiSeq|IlluminaNextSeq|IlluminaNovaSeq|IonProton|BGISeq",NA,NA,NA,NA,NA,"custom_enum",7,"02 Sequencing / QC"
"age","integer","non-unique","optional",FALSE,"Participant age in the unit specified by `age_unit`.","[0-9]+",NA,NA,NA,NA,NA,"integer",1,"03 Subject / Demographics"
"age_group","character","non-unique","optional",FALSE,"Age group classification: Newborn (< 1 month/28 days), Infant (>= 1 month & < 2 yrs), Child (>= 2 yrs & < 11 yrs), Adolescent (>= 11 yrs & < 18 yrs), Adult (>= 18 yrs & < 65 yrs), Elderly (>= 65 yrs) Static ontology IDs: NCIT:C16731; NCIT:C27956; NCIT:C16423; NCIT:C27954; NCIT:C17600; NCIT:C16268. Type: static.","Newborn|Infant|Child|Adolescent|Adult|Elderly","NCIT:C16731|NCIT:C27956|NCIT:C16423|NCIT:C27954|NCIT:C17600|NCIT:C16268",NA,NA,NA,NA,"static_enum",2,"03 Subject / Demographics"
"age_unit","character","non-unique","optional",FALSE,"Unit used to report `age`. Static ontology IDs: NCIT:C25301; NCIT:C29844; NCIT:C29846; NCIT:C29848. Type: static.","Day|Week|Month|Year","NCIT:C25301|NCIT:C29844|NCIT:C29846|NCIT:C29848",NA,NA,NA,NA,"static_enum",3,"03 Subject / Demographics"
"ancestry","character","non-unique","optional",FALSE,"Broad ancestry category (children of HANCESTRO:0004). Lookup seed ontology IDs: HANCESTRO:0004. Type: looked_up. Looked up: Yes. Includes child terms: Yes. Existing-dataset-only values now: Conditional (Yes only when lookup falls back to curated existing dataset values)",NA,NA,"HANCESTRO:0004","children",NA,NA,"dynamic_enum",4,"03 Subject / Demographics"
"ancestry_details","character","non-unique","optional",TRUE,"More specific ancestry term(s) typically refer to descendants of the selected ancestry category. Lookup seed ontology IDs: HANCESTRO:0004. Type: looked_up. Looked up: Yes. Includes child terms: Yes. Existing-dataset-only values now: Conditional (Yes only when lookup falls back to curated existing dataset values)",NA,NA,"HANCESTRO:0004","descendant",";",NA,"dynamic_enum",5,"03 Subject / Demographics"
"family","character","non-unique","optional",FALSE,"Family identifier when multiple related participants are included.",".+",NA,NA,NA,NA,NA,"any",6,"03 Subject / Demographics"
"family_role","character","non-unique","optional",FALSE,"Participant role within the family (for family-based sampling).","child|mother|father",NA,NA,NA,NA,NA,"custom_enum",7,"03 Subject / Demographics"
"sex","character","non-unique","optional",FALSE,"Biological sex of the subject Static ontology IDs: NCIT:C16576; NCIT:C20197. Type: static.","Female|Male","NCIT:C16576|NCIT:C20197",NA,NA,NA,NA,"static_enum",8,"03 Subject / Demographics"
"zigosity","character","non-unique","optional",FALSE,"Twin zygosity status.","monozygotic|dizygotic",NA,NA,NA,NA,NA,"custom_enum",9,"03 Subject / Demographics"
"age_min","double","non-unique","optional",FALSE,"Minimum possible age. For samples only with 'age_group' information, this represents the definition of a given age group. For specific age information, 'age_min' and 'age_max' are identical.","[0-9]+\.?[0-9]*",NA,NA,NA,NA,NA,"regexp",10,"03 Subject / Demographics"
"age_max","double","non-unique","optional",FALSE,"Maximum possible age. For samples only with 'age_group' information, this represents the definition of a given age group. For specific age information, 'age_min' and 'age_max' are identical.","[0-9]+\.?[0-9]*",NA,NA,NA,NA,NA,"regexp",11,"03 Subject / Demographics"
"age_years","double","non-unique","optional",FALSE,"Age in years","[0-9]+\.?[0-9]*",NA,NA,NA,NA,NA,"regexp",12,"03 Subject / Demographics"
"antibiotics_current_use","character","non-unique","optional",FALSE,"Whether the participant is currently using antibiotics.","Yes|No",NA,NA,NA,NA,NA,"binary",13,"04 Clinical / Intervention"
"bmi","double","non-unique","optional",FALSE,"Body mass index (BMI; EFO:0004340) calculated as weight (kg) divided by height squared (m²).","[0-9]+\.?[0-9]*",NA,NA,NA,NA,NA,"regexp",14,"04 Clinical / Intervention"
"dietary_restriction","character","non-unique","optional",TRUE,"Dietary regime (partial match to SNOMED:182922004 or SNOMED:162536008)","omnivore|low_gluten|high_gluten|gluten_free|vegetarian|vegan|pescatarian|high_fiber|low_fiber|paleo|equal_protein_fat_carbs|high_complex_carbs|lactose_intolerance|others",NA,NA,NA,";",NA,"custom_enum",15,"04 Clinical / Intervention"
"disease","character","non-unique","optional",TRUE,"Reported disease or condition term(s) for the participant; use Healthy when no target condition is detected. Lookup seed ontology IDs: NCIT:C7057; MONDO:0000001. Type: looked_up. Looked up: Yes. Includes child terms: Yes. Existing-dataset-only values now: Conditional (Yes only when lookup falls back to curated existing dataset values)",NA,"NCIT:C115935","NCIT:C7057;MONDO:0000001","descendant",";",NA,"dynamic_enum;static_enum",16,"04 Clinical / Intervention"
"disease_details","character","non-unique","optional",TRUE,"More specific disease or condition term(s) for the participant; use Healthy when appropriate. Lookup seed ontology IDs: NCIT:C7057; MONDO:0000001. Type: looked_up. Looked up: Yes. Includes child terms: Yes. Existing-dataset-only values now: Conditional (Yes only when lookup falls back to curated existing dataset values)",NA,"NCIT:C115935","NCIT:C7057;MONDO:0000001","descendant",";",NA,"dynamic_enum;static_enum",17,"04 Clinical / Intervention"
"disease_response_orr","character","non-unique","optional",FALSE,"Overall response rate (ORR; NCIT:C96613) indicator for treatment response.","Yes|No",NA,NA,NA,NA,NA,"binary",18,"04 Clinical / Intervention"
"disease_response_pfs","character","non-unique","optional",FALSE,"Progression Free Survival (PFS, EFO:0004920): Progression free survival is a measurement from a defined time point e.g. diagnosis and indicates that the disease did not progress i.e. tumours did not increase in size and new incidences did not occur. PFS is usually used in analyzing results of treatment for advanced disease.","Yes|No",NA,NA,NA,NA,NA,"binary",19,"04 Clinical / Intervention"
"disease_response_pfs_month","integer","non-unique","optional",FALSE,"Progression-free survival follow-up time in months.","[0-9]+",NA,NA,NA,NA,NA,"integer",20,"04 Clinical / Intervention"
"disease_response_recist","character","non-unique","optional",FALSE,"Response Evaluation Criteria in Solid Tumors (RECIST, DICOM:112022): Standard parameters to be used when documenting response of solid tumors to treatment; a set of published rules that define when cancer patients improve (`respond`), stay the same (`stable`), or worsen (`progression`) during treatments. (from www.recist.com) Static ontology IDs: NCIT:C159715; NCIT:C159547; NCIT:C159716; NCIT:C159546. Type: static.","RECIST Complete Response|RECIST Partial Response|RECIST Progressive Disease|RECIST Stable Disease","NCIT:C159715|NCIT:C159547|NCIT:C159716|NCIT:C159546",NA,NA,NA,NA,"static_enum",21,"04 Clinical / Intervention"
"feces_phenotype","character","non-unique","optional",TRUE,"Stool-related phenotype or clinical measurement type. Static ontology IDs: SNOMED:443172007; NCIT:C82005; NCIT:C191036. Type: static.","Bristol stool form score (observable entity)|Calprotectin Measurement|Harvey-Bradshaw Index Clinical Classification","SNOMED:443172007|NCIT:C82005|NCIT:C191036",NA,NA,"<;>",NA,"static_enum",22,"04 Clinical / Intervention"
"fmt_id","character","non-unique","optional",TRUE,"Study-specific identifier assigned to FMT participants.",".+",NA,NA,NA,";",NA,"any",23,"04 Clinical / Intervention"
"fmt_role","character","non-unique","optional",FALSE,"Role in fecal microbiota transplantation (donor or recipient before/after procedure).","Recipient (after procedure)|Recipient (before procedure)|Donor",NA,NA,NA,NA,NA,"custom_enum",24,"04 Clinical / Intervention"
"hla","character","non-unique","optional",TRUE,"Human leukocyte antigen (HLA) typing information. Lookup seed ontology IDs: MRO:0001676. Type: looked_up. Looked up: Yes. Includes child terms: Yes. Existing-dataset-only values now: Conditional (Yes only when lookup falls back to curated existing dataset values)",NA,NA,"MRO:0001676","descendant",";",NA,"dynamic_enum",25,"04 Clinical / Intervention"
"neonatal_birth_weight","numeric","non-unique","optional",FALSE,"Birth weight (EFO:0004344) recorded in grams.","^[1-9]\d*(\.\d+)?$",NA,NA,NA,NA,NA,"numeric",26,"04 Clinical / Intervention"
"neonatal_delivery_procedure","character","non-unique","optional",FALSE,"Delivery method at birth (NCIT:C81179). Static ontology IDs: NCIT:C114141; NCIT:C92772; NCIT:C46088; NCIT:C81303. Type: static.","Elective Cesarean Delivery|Emergency Cesarean Delivery|Cesarean Section|Vaginal Delivery","NCIT:C114141|NCIT:C92772|NCIT:C46088|NCIT:C81303",NA,NA,NA,NA,"static_enum",27,"04 Clinical / Intervention"
"neonatal_feeding_method","character","non-unique","optional",TRUE,"Infant feeding method(s) during early life.","Mixed Feeding|Exclusively Breastfeeding|Exclusively Formula Feeding|No Breastfeeding",NA,NA,NA,";",NA,"custom_enum",28,"04 Clinical / Intervention"
"neonatal_gestational_age","numeric","non-unique","optional",FALSE,"Gestational age at birth (EFO:0005112) in weeks.","^[1-9]\d*(\.\d+)?$",NA,NA,NA,NA,NA,"numeric",29,"04 Clinical / Intervention"
"neonatal_preterm_birth","character","non-unique","optional",FALSE,"Birth when a fetus is less than 37 weeks and 0 days gestational age (NCIT:C92861). Static ontology IDs: NCIT:C92861; NCIT:C114093. Type: static.","Preterm Birth|Term Birth","NCIT:C92861|NCIT:C114093",NA,NA,NA,NA,"static_enum",30,"04 Clinical / Intervention"
"obgyn_birth_control","character","non-unique","optional",FALSE,"Whether oral contraceptive birth control is currently used.","Yes|No",NA,NA,NA,NA,NA,"binary",31,"04 Clinical / Intervention"
"obgyn_lactating","character","non-unique","optional",FALSE,"An indication that the subject is currently producing milk. (NCIT:C82463)","Yes|No",NA,NA,NA,NA,NA,"binary",32,"04 Clinical / Intervention"
"obgyn_menopause","character","non-unique","optional",FALSE,"Menopausal status of the participant. Lookup seed ontology IDs: NCIT:C106541. Type: looked_up. Looked up: Yes. Includes child terms: Yes. Existing-dataset-only values now: Conditional (Yes only when lookup falls back to curated existing dataset values)",NA,NA,"NCIT:C106541","descendant",NA,NA,"dynamic_enum",33,"04 Clinical / Intervention"
"obgyn_pregnancy","character","non-unique","optional",FALSE,"Current pregnancy status of the participant. Static ontology IDs: NCIT:C124295; NCIT:C82475. Type: static.","Pregnant|Not Pregnant","NCIT:C124295|NCIT:C82475",NA,NA,NA,NA,"static_enum",34,"04 Clinical / Intervention"
"smoker","character","non-unique","optional",TRUE,"Tobacco smoking status/history; use `Non-smoker (finding)` only when past smoking history is unavailable. Static ontology IDs: SNOMED:77176002; SNOMED:8392000; SNOMED:8517006; SNOMED:266919005. Type: static.","Smoker (finding)|Non-smoker (finding)|Ex-smoker (finding)|Never smoked tobacco (finding)","SNOMED:77176002|SNOMED:8392000|SNOMED:8517006|SNOMED:266919005",NA,NA,";",NA,"static_enum",35,"04 Clinical / Intervention"
"treatment","character","non-unique","optional",TRUE,"Treatment(s) or medication(s) administered to the participant. Static ontology IDs: NCIT:C41132. Lookup seed ontology IDs: NCIT:C1908. Type: static + looked_up. Looked up: Yes. Includes child terms: Yes. Existing-dataset-only values now: Conditional (Yes only when lookup falls back to curated existing dataset values)",NA,"NCIT:C41132","NCIT:C1908","descendant",";",NA,"dynamic_enum;static_enum",36,"04 Clinical / Intervention"
"tumor_staging_ajcc","character","non-unique","optional",FALSE,"American Joint Committee on Cancer (tumor staging) (SNOMED:258236004). A system to describe the amount and spread of cancer in a patient's body","0|I|II|III|IV|I/II|III/IV|IIIA|IIIB",NA,NA,NA,NA,NA,"static_enum",37,"04 Clinical / Intervention"
"tumor_staging_tnm","character","non-unique","optional",FALSE,"Tumor-node-metastasis (TNM) tumor staging system (tumor staging) (SNOMED:254293002).A system to describe the amount and spread of cancer in a patient's body","T[X1-4]N[X0-3]M[X0-1]|pTis|Tis",NA,NA,NA,NA,NA,"regexp",38,"04 Clinical / Intervention"
"biomarker_name","character","non-unique","optional",TRUE,"A measurable and quantifiable characteristic or substance that serves as an indicator of a biological state, condition, or process within an organism.",NA,NA,NA,NA,";",NA,"dynamic_enum",39,"04 Clinical / Intervention"
"biomarker_unit","character","non-unique","optional",TRUE,"Unit for biomarker",NA,NA,NA,NA,";",NA,"any",40,"04 Clinical / Intervention"
"biomarker_value","numeric","non-unique","optional",TRUE,"Value for biomarker","[0-9]+\.?[0-9]*",NA,NA,NA,";",NA,"regexp",41,"04 Clinical / Intervention"
"disease_response_os","numeric","non-unique","optional",FALSE,"Overall survival duration (time to death from any cause; NCIT:C125201).","^[1-9]\d*(\.\d+)?$",NA,NA,NA,NA,NA,"numeric",42,"04 Clinical / Intervention"
"disease_response_os_unit","character","non-unique","optional",FALSE,"Unit used to report `disease_response_os`. Static ontology IDs: NCIT:C25301; NCIT:C29844; NCIT:C29846; NCIT:C29848. Type: static.","Day|Week|Month|Year","NCIT:C25301|NCIT:C29844|NCIT:C29846|NCIT:C29848",NA,NA,NA,NA,"static_enum",43,"04 Clinical / Intervention"
"ecog_performance_status","character","non-unique","optional",FALSE,"A performance status scale designed to assess disease progression and its affect on the daily living abilities of the patient. (NCIT:C105721) Static ontology IDs: NCIT:C105722; NCIT:C105723; NCIT:C105724; NCIT:C105725; NCIT:C105726; NCIT:C105727; NCIT:C105728. Type: static.","ECOG Performance Status 0|ECOG Performance Status 1|ECOG Performance Status 2|ECOG Performance Status 2 or Higher|ECOG Performance Status 3|ECOG Performance Status 4|ECOG Performance Status 5","NCIT:C105722|NCIT:C105723|NCIT:C105724|NCIT:C105725|NCIT:C105726|NCIT:C105727|NCIT:C105728",NA,NA,NA,NA,"static_enum",44,"04 Clinical / Intervention"
"feces_phenotype_value","character","non-unique","optional",TRUE,"Value(s) corresponding to `feces_phenotype` measurements.","^\d+(\.\d+)?$",NA,NA,NA,"<;>",NA,"character",45,"04 Clinical / Intervention"
"tumor_size_measurement","numeric","non-unique","optional",FALSE,"Tumor size measurement from clinical assessment or resected specimen (NCIT:C106303).","^[1-9]\d*(\.\d+)?$",NA,NA,NA,NA,NA,"numeric",46,"04 Clinical / Intervention"
"tumor_size_residual_measurement","character","non-unique","optional",FALSE,"Residual tumor size measurement after treatment/procedure (NCIT:C198194).","^[1-9]\d*(\.\d+)?$",NA,NA,NA,NA,NA,"numeric",47,"04 Clinical / Intervention"
"uncurated_metadata","character","non-unique","optional",TRUE,"Additional free-text metadata not represented by existing fields.",".+",NA,NA,NA,"<;>",NA,"any",1,"05 Other"
"host_species","character","non-unique","required",FALSE,"Host species of the participant.",NA,NA,"NCBITaxon:1","descendant",NA,NA,"dynamic_enum",14,"01 Sampling / Study"
"""

_default_schema_cache = None


def _parse_schema_csv_text(csv_text: str) -> dict:
    """Parse a data-dictionary/codebook CSV's text into the same schema-dict shape
    additional_pipeline.fetch_standardization_schema() builds from a live CSV URL:
    {field_name: {"description": str, "allowed_values": list, "required": bool}}.
    Column-detection logic mirrors that function (kept in sync manually since
    model.py must not import additional_pipeline -- see circular-import note at
    additional_pipeline.py's own `model = _try_import("model")`).
    """
    schema: dict = {}
    lines = csv_text.splitlines()
    reader = csv.DictReader(lines)
    headers = reader.fieldnames or []
    if not headers:
        return schema

    name_col = next(
        (h for h in headers if any(k in h.lower() for k in ("name", "field", "variable", "column"))),
        headers[0]
    )
    desc_col = next(
        (h for h in headers if any(k in h.lower() for k in ("description", "definition", "label", "detail"))),
        headers[1] if len(headers) > 1 else None
    )
    # Prefer a header explicitly saying "allowed" (e.g. "allowedvalues") over the
    # broader value/code/category/option keywords -- avoids matching a
    # "multiplevalues" boolean-flag column (whether the field permits multiple
    # values) ahead of the actual allowed-values column.
    val_col = next((h for h in headers if "allowed" in h.lower()), None)
    if not val_col:
        val_col = next(
            (h for h in headers
             if any(k in h.lower() for k in ("value", "code", "category", "option"))
             and "multiplevalue" not in h.lower().replace("_", "").replace(" ", "")),
            None
        )
    required_col = next((h for h in headers if "required" in h.lower()), None)
    is_codebook = val_col is not None

    for row in reader:
        field = (row.get(name_col) or "").strip()
        if not field:
            continue
        desc = (row.get(desc_col) or "").strip() if desc_col else ""
        val = (row.get(val_col) or "").strip() if val_col else ""
        req_raw = (row.get(required_col) or "").strip().lower() if required_col else ""
        is_required = req_raw in ("required", "true", "yes", "1")

        if field not in schema:
            schema[field] = {"description": desc, "allowed_values": [], "required": is_required}
        else:
            if desc and not schema[field]["description"]:
                schema[field]["description"] = desc
            if is_required:
                schema[field]["required"] = True

        if is_codebook and val and val.upper() not in ("NA", "N/A"):
            for v in val.split("|"):
                v = v.strip()
                if v and v not in schema[field]["allowed_values"]:
                    schema[field]["allowed_values"].append(v)

    return schema


def _get_default_schema() -> dict:
    """Memoized parse of the built-in default schema (parsed once, reused)."""
    global _default_schema_cache
    if _default_schema_cache is None:
        _default_schema_cache = _parse_schema_csv_text(_DEFAULT_SCHEMA_CSV_TEXT)
    return _default_schema_cache


def _build_schema_hint(fields, schema: dict) -> str:
    """Build the STANDARDIZATION RULES prompt block for the given fields from a
    schema dict {field: {"description": str, "allowed_values": list}}. Shared by
    Pass 1 (multi_prompts) and Pass 2 (_extract_additional_fields) so both use
    identical field-definition/allowed-value/CONTROL-boolean phrasing. General --
    works with any schema dict (a user-supplied CSV or the built-in default),
    not tied to any particular set of field names.
    """
    if not fields or not schema:
        return ""
    schema_lines = []
    for f in fields:
        entry = schema.get(f, {})
        if isinstance(entry, dict):
            desc = entry.get("description", "")
            allowed = entry.get("allowed_values", [])
        else:
            desc = str(entry)
            allowed = []
        if not desc and not allowed:
            continue  # field not in this schema -- nothing to annotate
        line = f"  - {f}"
        if desc:
            line += f": {desc}"
        # Any field whose name is/contains "control" gets an explicit definition,
        # regardless of whether its allowed values look boolean or enum-shaped --
        # different schemas encode case/control differently (TRUE/FALSE, or
        # "Study Control"/"Case"/..., or free text), but the underlying concept
        # and the risk of the model guessing "case" by default is the same.
        _is_control_field = bool(re.search(r'\bcontrol\b', f.lower()))
        _control_def = (
            "CONTROL means the sample belongs to the group with NONE of the "
            "study's conditions/exposures present -- the fully unaffected/"
            "reference group -- not merely 'not the primary condition being "
            "studied' (a sample with a different or secondary condition is NOT "
            "a control). If you cannot confidently determine full-unaffected "
            "status for this specific sample, output 'unknown' rather than "
            "defaulting to a case/disease label."
        ) if _is_control_field else ""
        if allowed:
            bool_vals = {v.strip().lower() for v in allowed}
            if bool_vals <= {"true", "false", "0", "1", "yes", "no"}:
                line += (
                    f" [BOOLEAN — output TRUE if sample IS a control/reference, "
                    f"FALSE if sample is a case/disease/treatment group. {_control_def} "
                    f"Allowed: {', '.join(str(v) for v in allowed[:10])}. "
                    f"Do NOT output 'unknown' when you can confidently classify case vs. control.]"
                )
            else:
                line += f" [allowed values: {', '.join(str(v) for v in allowed[:20])}. {_control_def}]"
        elif _is_control_field:
            line += f" [{_control_def}]"
        schema_lines.append(line)

    if not schema_lines:
        return ""
    return (
        "STANDARDIZATION RULES — use these exact field definitions and "
        "allowed values from the schema:\n"
        + "\n".join(schema_lines)
        + "\n\nIMPORTANT: Use ONLY the allowed values listed above. "
        "Choose the closest match when exact wording differs. "
        "Write 'unknown' ONLY when the information is genuinely absent from ALL source texts.\n"
    )


def multi_prompts(dictsAccs, output_format_str, niche_cases=None, prompt_template="default",
                  standardization_schema=None):
  """Build per-accession prompts.

  standardization_schema: dict {field_name: {"description": str, "allowed_values": list}}
  from a schema CSV (e.g. cMD data dictionary + codebook). When omitted/empty, falls
  back to the built-in default schema (_get_default_schema()) so every run gets some
  standardization guidance even without a user-supplied schema URL. When a real dict
  is supplied (e.g. fetched from a user's own standardization_url), it's used instead.
  Each requested field is annotated with its schema definition AND allowed values so
  the LLM constrains its output to the canonical vocabulary.
  """
  prompts = {}
  if niche_cases:
    fields_list = ", ".join(niche_cases)
    _schema = standardization_schema if standardization_schema else _get_default_schema()
    schema_hint = _build_schema_hint(niche_cases, _schema)

    # Build per-sample-table hint if disease/control fields are requested
    _field_lower = [f.lower() for f in niche_cases]
    _needs_disease_hint = any(
        kw in " ".join(_field_lower)
        for kw in ("disease", "control", "phenotype", "condition", "health", "group", "status", "diagnosis")
    )
    _disease_hint = (
        "CONTROL DEFINITION: a 'control' sample belongs to the group with NONE of the study's conditions/exposures present (the fully unaffected/reference group) -- "
        "not merely 'not the primary condition being studied.' If you cannot confidently determine full-unaffected status for this sample, output 'unknown' rather than "
        "defaulting to a case/disease label.\n"
        "Target condition definition: Primary phenotype(s), condition(s), or disease status THAT THIS SPECIFIC SAMPLE ACTUALLY HAS in the study, "
        "as determined by this sample's group assignment or individual clinical status. NOT the study's overall research topic.\n"
        "CROSS-FIELD CONSISTENCY: several requested fields may describe the same underlying case/control assignment from different angles "
        "(e.g. a case/control-group field alongside a condition/phenotype/diagnosis/health-status field). Whatever you conclude for one such "
        "field, every other such field MUST agree for THIS SAME sample -- never let one field say this sample is a control/unaffected while "
        "another field names a condition as if this sample has it, or vice versa. If this sample is the control/unaffected group, every "
        "condition/phenotype/diagnosis-type field must explicitly reflect that absence (e.g. 'none', 'not applicable', 'unaffected', "
        "'no <condition>') rather than restating whatever condition(s) the study investigates overall. If this sample IS affected/exposed, "
        "name only the specific condition(s)/phenotype(s) assigned to THIS sample individually, never the full list of conditions the study "
        "investigates as a whole."
    ) if _needs_disease_hint else ""

    # Build a hint for study/dataset-identifier fields so the model doesn't
    # accept an accession number as the answer (an accession is not a study name).
    _needs_study_name_hint = any(
        kw in " ".join(_field_lower)
        for kw in ("study_name", "study name", "studyname", "dataset_name")
    )
    _study_name_hint = (
        "IMPORTANT — for study/dataset-identifier fields: use the paper's own identifying "
        "convention (e.g. first-author surname + publication year, such as 'SmithJ_2021'), or "
        "the actual publication title if no such convention is stated. Do NOT output an "
        "NCBI/ENA accession number (BioProject, BioSample, SRA/ENA study ID, run accession, "
        "etc.) as the answer -- an accession-shaped string is never a valid study/dataset name.\n"
    ) if _needs_study_name_hint else ""

    _caveat = (
    "IMPORTANT — per-subject assignment: many studies have multiple participant groups "
    "(e.g. a reference/control group and one or more affected/exposed groups). "
    "Search for a table or supplementary file that maps individual sample identifiers (sample IDs / subject IDs / NCBI BioSample, isolate name, etc. in the NCBI records) "
    "accessions to their specific group. The accession being analysed is shown in 'Prompt N:' above — find its row in that table and extract the group/condition for "
    "THAT SPECIFIC SAMPLE, not the study as a whole. "
    "PRIORITY RULE: a table/section that maps individual sample identifiers to their specific category is stronger evidence than general prose describing the study's groups or conditions as a whole; "
    "when such a table/section exist, cite its matching row, not the prose. "
    "COMPLETENESS: do not stop at the first candidate table/section you find. First check what identifying attribute(s) actually exist on this sample's own record, "
    "whatever they are called (id, subject_id, isolate_name, strain, specimen_code, patient number, or anything else) -- "
    "then check EVERY table/section in the source text that uses a matching or clearly related identifying scheme, "
    "not just the first or most prominent one. Check NCBI BioSample attributes and also paper tables / supplementary metadata tables. "
    "Do NOT report the full list of study groups — report only the group for this individual sample. "
    ) if _needs_disease_hint else ""

    niche_prompt = (
      f"Extract the following metadata fields: {fields_list}.\n"
      f"For EACH FIELD: Extract the SAMPLE'S ACTUAL STATUS or INDIVIDUAL VALUE (what applies to THIS sample), "
      f"never the study's general description or objectives. Even for fields defined in terms of 'the study,' prioritize this sample's specific measured/assigned value. "
      f"Only use study-level information as fallback when this sample has no individual value recorded.\n"
      f"{schema_hint}"
      f"{_disease_hint}"
      f"{_study_name_hint}"
      f"{_caveat}"
      f"For each field: find the most specific value stated anywhere in the sources.\n"
      f"Infer from context when not explicit. Write 'unknown' ONLY when truly absent from all sources.\n"
      f"If different sources give DIFFERENT values for the same field, keep the most specific/reliable\n"
      f"value AND append '##CONFLICT: source_A=<val_A>, source_B=<val_B>' to that value so the conflict is visible.\n"
    )
  else:
    niche_prompt = ""

  for acc_pos in range(len(list(dictsAccs.keys()))):
    acc = list(dictsAccs.keys())[acc_pos]
    acc_cleaned = acc.split('.')[0] if acc else acc
    accession_found_in_text = False
    context_for_llm = dictsAccs[acc]
    if prompt_template == "default":
      field_count = len(output_format_str.split(", "))
      _fmt_fields = [f.strip() for f in output_format_str.split(",")]
      _has_country_fmt = any('country' in f.lower() for f in _fmt_fields)
      _has_type_fmt = any('modern' in f.lower() or 'ancient' in f.lower() for f in _fmt_fields)
      _country_hint = (
          "Identify its primary associated geographic location (country preferred; "
          "fall back to region/continent if no country mentioned; write 'unknown' "
          "if no geographic clues are present).\n"
      ) if _has_country_fmt else ""
      _type_hint = (
          "Determine if the sample source is 'modern' (living individual) or "
          "'ancient' (prehistoric/archaeological); assume 'modern' if not specified.\n"
      ) if _has_type_fmt else ""
      prompt_for_llm = (
      f"Prompt {acc_pos+1}: "
      f"Given the following text snippets, analyze the biological sample with "
      f"accession number {acc_cleaned}.\n"
      f"{_country_hint}"
      f"{_type_hint}"
      f"{niche_prompt}"
      f"\nOUTPUT FORMAT (follow exactly):\n"
      f"This may be one of SEVERAL numbered prompts in this same request, each analyzing a DIFFERENT "
      f"accession. Before writing this prompt's field blocks, output exactly this header line on its own, "
      f"with no other text on it:\n"
      f"=== ANSWERS FOR PROMPT {acc_pos+1} (accession {acc_cleaned}) ===\n"
      f"Then write ONE block per field below, in EXACTLY this order ({field_count} fields): {output_format_str}\n"
      f"Each block MUST have this exact 3-line structure:\n"
      f"FIELD: <field_name>\n"
      f"REASONING: <one narrative sentence citing exactly WHERE — name the specific table/section/figure "
      f"and include a verbatim excerpt (≤15 words) that confirms the value> "
      f"[Sources: <key1> (<location>, '<verbatim excerpt>'); <key2> (<location>, '<verbatim excerpt>')] "
      f"[Conflict: <describe any disagreement between sources, or 'none'>] "
      f"[ID-match: true|false]\n"
      f"ANSWER: <the value for this field>\n"
      f"  - key in [Sources: ...] = exact header from 'The source - <key>:' blocks in the text "
      f"(e.g. NCBI_biosample, a URL, user_uploaded_file); list EVERY source that mentions this field.\n"
      f"  - [ID-match: true] ONLY if this value was confirmed by finding this sample's own numeric ID/identifier "
      f"as a row in a numbered table (per the PRIORITY RULE above, when applicable); false if the value instead "
      f"came from general prose or a topical description not tied to this sample's specific ID. For fields read "
      f"directly from this sample's own BioSample/SRA record (no table lookup needed), output true.\n"
      f"  - CRITICAL: write REASONING first, decide ANSWER only AFTER, based on that reasoning — never write "
      f"ANSWER before you have worked out the REASONING for that same field. ANSWER must agree with its own "
      f"REASONING and with every other field's ANSWER for this sample (see CROSS-FIELD CONSISTENCY above).\n"
      f"  - For conflicting values across sources, in ANSWER write: <best_value> ##CONFLICT: source_A=<val_A>, source_B=<val_B>\n"
      f"  - Separate each field's block from the next with a blank line. No markdown, no extra headers, no summary line.\n"
      f"Example for 2 fields (country_name, modern/ancient) — placeholder values only:\n"
      f"FIELD: country_name\n"
      f"REASONING: The geo_loc_name attribute is 'Spain: Region' and subjects were recruited locally. [Sources: NCBI_biosample (geo_loc_name attribute, 'Spain: Region'); https://doi.org/10.9999/example (Methods section, 'recruited in Region')] [Conflict: none] [ID-match: true]\n"
      f"ANSWER: Spain\n"
      f"\n"
      f"FIELD: modern/ancient\n"
      f"REASONING: Subjects were described as living participants enrolled in 2020 per the Methods section. [Sources: https://doi.org/10.9999/example (Methods section, 'living participants enrolled in 2020')] [Conflict: none] [ID-match: false]\n"
      f"ANSWER: modern\n"
      f"\nText Snippets:\n{context_for_llm}")
      if acc_cleaned.lower() in context_for_llm.lower():
        accession_found_in_text = True
      prompts[acc] = [prompt_for_llm, accession_found_in_text]
  return prompts

def standardize_with_llm(extracted_values: dict, schema: dict, acc: str) -> dict:
    """
    Post-extraction LLM standardization.
    Maps extracted free-text values to canonical schema-defined values.
    Uses Anthropic Claude first (better instruction-following), Gemini as fallback.
    """
    if not schema or not extracted_values:
        return extracted_values

    # Pop raw schema text (for prompt context) before field lookup
    schema_text = schema.pop('__schema_text__', '') or ''

    # Only standardize fields that exist in the schema
    fields_to_std = {k: v for k, v in extracted_values.items()
                     if k in schema and v and v.lower() != "unknown"}
    if not fields_to_std:
        if schema_text:
            schema['__schema_text__'] = schema_text  # restore for next call
        return extracted_values

    schema_lines = []
    for field, val in fields_to_std.items():
        entry = schema.get(field, {})
        if isinstance(entry, dict):
            desc = entry.get("description", "")
            allowed = entry.get("allowed_values", [])
        else:
            desc = str(entry)
            allowed = []
        line = f"  {field}: current='{val}'"
        if desc:
            line += f", definition='{desc}'"
        if allowed:
            bool_vals = {v.strip().lower() for v in allowed}
            if bool_vals <= {"true", "false", "0", "1", "yes", "no"}:
                line += (f", BOOLEAN — TRUE=is a control, FALSE=is a case/disease. "
                         f"Allowed: {', '.join(str(v) for v in allowed[:10])}")
            else:
                line += f", allowed=[{', '.join(str(v) for v in allowed[:15])}]"
        schema_lines.append(line)

    schema_context = (
        f"REFERENCE SCHEMA (use this to understand field definitions and allowed values):\n"
        f"{schema_text[:5000]}\n\n"
        if schema_text else ""
    )

    prompt = (
        f"You are a biomedical metadata standardizer for sample {acc}.\n\n"
        f"{schema_context}"
        f"Map each extracted value to its canonical schema value:\n"
        + "\n".join(schema_lines) + "\n\n"
        "Rules:\n"
        "1. Use ONLY the allowed values listed; pick the closest match.\n"
        "2. For BOOLEAN fields: if the sample is clearly a case/disease/treatment, output FALSE for 'control'. "
        "If it is clearly in the control/reference group, output TRUE.\n"
        "3. If you cannot determine the correct standardized value, keep the original.\n"
        "4. Return ONLY a JSON object: {\"field\": \"standardized_value\"}.\n"
        "No markdown, no explanation.\n"
    )
    if schema_text:
        schema['__schema_text__'] = schema_text  # restore for next call

    try:
        response_text, _ = call_llm_api(prompt)
        raw = response_text.strip()
        if raw.startswith("```"):
            parts = raw.split("```")
            raw = parts[1] if len(parts) > 1 else raw
            if raw.startswith("json"):
                raw = raw[4:]
        result = json.loads(raw.strip())
        standardized = dict(extracted_values)
        for k, v in result.items():
            if k in extracted_values and v is not None:
                standardized[k] = str(v).strip()
        return standardized
    except Exception as e:
        print(f"[standardize_with_llm] WARNING: {e}")
        return extracted_values


def align_to_schema(extracted_dict: dict, schema: dict, acc: str) -> dict:
    """
    Map free-text extracted fields to the closest matching schema field names
    and standardize their values to the schema's allowed values.

    Returns {schema_field: {"value": standardized_value, "from_field": original_key}}
    only for high-confidence matches. "from_field" lets the caller remove the
    duplicate raw-named entry once it's been folded into the canonical name.
    """
    if not schema or not extracted_dict:
        return {}
    schema_fields = [k for k in schema if not k.startswith('__')]
    if not schema_fields:
        return {}

    schema_lines = []
    for f in schema_fields[:40]:
        entry = schema.get(f, {})
        if isinstance(entry, dict):
            desc = entry.get('description', '')
            allowed = entry.get('allowed_values', [])
        else:
            desc = str(entry)
            allowed = []
        line = f"  {f}"
        if desc:
            line += f": {desc}"
        if allowed:
            bool_vals = {v.strip().lower() for v in allowed}
            if bool_vals <= {"true", "false", "0", "1", "yes", "no"}:
                line += (f" [BOOLEAN: TRUE=control/reference, FALSE=case/disease. "
                         f"Allowed: {', '.join(str(v) for v in allowed[:6])}]")
            else:
                line += f" [allowed: {', '.join(str(v) for v in allowed[:12])}]"
        schema_lines.append(line)

    ext_lines = "\n".join(f"  {k}: {v}" for k, v in list(extracted_dict.items())[:60])

    prompt = (
        f"Sample accession: {acc}\n\n"
        "EXTRACTED METADATA (free-text field names and values):\n"
        f"{ext_lines}\n\n"
        "TARGET SCHEMA (canonical field names + allowed values):\n"
        + "\n".join(schema_lines) + "\n\n"
        "Task: For each extracted field, find the best-matching schema field name.\n"
        "Standardize the value to the nearest allowed value when one is listed.\n"
        "Rules:\n"
        "1. Only include fields where you're confident of the match.\n"
        "2. For BOOLEAN fields: if the sample is clearly a case/disease/treatment → FALSE. "
        "If clearly a control/reference → TRUE.\n"
        "3. Never invent values not in the allowed list.\n"
        "4. If a value matches an allowed entry regardless of case/punctuation, use the "
        "exact allowed value string.\n"
        "Return ONLY a JSON object: {\"schema_field_name\": {\"value\": \"standardized_value\", "
        "\"from_field\": \"<the original extracted field name above that this came from>\"}}.\n"
        "No markdown fences, no explanation.\n"
    )

    try:
        response_text, _ = call_llm_api(prompt)
        raw = response_text.strip()
        if raw.startswith('```'):
            parts = raw.split('```')
            raw = parts[1] if len(parts) > 1 else raw
            if raw.lower().startswith('json'):
                raw = raw[4:]
        brace_start = raw.find('{')
        brace_end   = raw.rfind('}')
        if brace_start != -1 and brace_end > brace_start:
            raw = raw[brace_start:brace_end + 1]
        result = json.loads(raw.strip())
        out: dict = {}
        for k, v in result.items():
            if k not in schema or v is None:
                continue
            if isinstance(v, dict):
                val = str(v.get('value', '')).strip()
                from_field = str(v.get('from_field', '')).strip()
            else:
                val = str(v).strip()
                from_field = ''
            if val:
                out[k] = {'value': val, 'from_field': from_field}
        return out
    except Exception as e:
        print(f"[align_to_schema] WARNING: {e}")
        return {}


def annotate_with_ontologies(extracted_dict: dict, context_text: str, acc: str) -> dict:
    """
    Group extracted metadata into 5 biological ontology categories and assign
    standard ontology IDs (NCBITaxon, UBERON, OBI, CHMO, MS, GO, DOID, PATO, SO).

    Returns dict with 5 category keys, each a list of 'ONTOLOGY:ID | label' strings.
    """
    if not extracted_dict and not context_text:
        return {}

    extracted_str = "\n".join(
        f"  {k}: {v}" for k, v in list(extracted_dict.items())[:40]
    )
    ctx_snippet = context_text[:10000] if len(context_text) > 10000 else context_text

    prompt = (
        f"Sample: {acc}\n\n"
        "Extracted metadata:\n"
        f"{extracted_str}\n\n"
        "Source text excerpt:\n"
        f"{ctx_snippet}\n\n"
        "Task: Annotate this biological sample metadata with standard ontology IDs.\n\n"
        "Use IDs from these ontologies:\n"
        "  NCBITaxon — organism taxonomy (e.g. NCBITaxon:10090 | Mus musculus)\n"
        "  UBERON    — anatomical parts / tissues / body sites\n"
        "  OBI       — assay / instrument / sample preparation types\n"
        "  CHMO      — chemical methods / chromatography\n"
        "  MS        — mass spectrometry terms (Proteomics Standards Initiative)\n"
        "  GO        — biological processes / molecular functions / cellular components\n"
        "  DOID      — human diseases\n"
        "  PATO      — phenotypic qualities (age, sex, genotype)\n"
        "  SO        — sequence ontology (genotype, mutation type)\n\n"
        "Assign terms to exactly these 5 categories:\n"
        "  taxonomy              — organism species / strain\n"
        "  organism_part         — body site, tissue, organ, anatomical region\n"
        "  host_characteristics  — sex, age, disease status, genotype, phenotype\n"
        "  experimental_conditions — assay type, instrument, extraction method, data format\n"
        "  contextual_study      — biological process, study design, disease context\n\n"
        "Format each entry as 'ONTOLOGY:ID | label'.\n"
        "Include multiple entries per category when warranted (one per line in the list).\n"
        "Return ONLY valid JSON with exactly these 5 keys:\n"
        "{\n"
        '  "taxonomy": ["NCBITaxon:10090 | Mus musculus"],\n'
        '  "organism_part": ["UBERON:0006909 | lumen of digestive tract"],\n'
        '  "host_characteristics": ["PATO:0000384 | male", "DOID:9352 | type 2 diabetes"],\n'
        '  "experimental_conditions": ["OBI:0000470 | mass spectrometry assay"],\n'
        '  "contextual_study": ["GO:0006955 | immune response"]\n'
        "}\n"
        "No markdown fences. Return only valid JSON.\n"
    )

    try:
        response_text, _ = call_llm_api(prompt)
        raw = response_text.strip()
        if raw.startswith('```'):
            parts = raw.split('```')
            raw = parts[1] if len(parts) > 1 else raw
            if raw.lower().startswith('json'):
                raw = raw[4:]
        brace_start = raw.find('{')
        brace_end   = raw.rfind('}')
        if brace_start != -1 and brace_end > brace_start:
            raw = raw[brace_start:brace_end + 1]
        result = json.loads(raw.strip())
        expected = {'taxonomy', 'organism_part', 'host_characteristics',
                    'experimental_conditions', 'contextual_study'}
        return {k: v for k, v in result.items() if k in expected}
    except Exception as e:
        print(f"[annotate_with_ontologies] WARNING: {e}")
        return {}


async def getMoreInfoForAcc(iso=None, acc=None, saveLinkFolder=None, niche_cases=None, limit_context=250000, extra_metadata=None):
  linksWithTexts, links, context_for_llm = {}, [], ""
  meta_expand = smart_fallback.fetch_ncbi(acc)
  if extra_metadata:
    for k, v in extra_metadata.items():
      if v and str(v).lower() not in ("unknown", ""):
        meta_expand[k] = v
  raw_tem_links = smart_fallback.smart_google_search(acc, meta_expand)
  tem_links = pipeline.unique_preserve_order(raw_tem_links)
  print("this is tem links with acc: ", tem_links)
  # filter the quality link
  print("start the smart filter link")
  #success_process, output_process = run_with_timeout(smart_fallback.filter_links_by_metadata,args=(tem_links,saveLinkFolder),kwargs={"accession":acc},timeout=90)
  output_process = await smart_fallback.async_filter_links_by_metadata(
      tem_links, saveLinkFolder, accession=acc
  )
  print('inside getMoreInfoForAcc and here is outputProcess: ', output_process)
  if output_process:
    linksWithTexts.update(output_process)
    print("yeah we have linksWithTexts and len: ", len(linksWithTexts))
    print("yes succeed for smart filter link")
    links += list(linksWithTexts.keys())
    print("link keys: ", links)
  else: 
    print("not have output_process")
    links += tem_links      
  if links:
    # use build context for llm function to reduce token
    texts_reduce = []
    linksWithTexts_reduce = {}
    reduce_context_for_llm = ""
    print("links:", links)
    for link in links:
      print("link: ", link)
      new_all_output = await pipeline.process_link_allOutput(link, 
                iso, acc, saveLinkFolder, linksWithTexts_reduce, context_for_llm)
      print("done all output")
      context_for_llm += new_all_output
      texts_reduce.append(new_all_output)
      linksWithTexts_reduce[link] = {"all_output": new_all_output}
    # tasks = [
    #     pipeline.process_link_allOutput(link, iso, acc, saveLinkFolder, linksWithTexts, all_output)
    #     for link in links
    # ]
    # results = await asyncio.gather(*tasks)
    # print("this is result:", results)
    # # combine results
    # for new_all_output in results:
    #   context_for_llm += new_all_output
    print("len of context after merge all: ", len(context_for_llm))

  if len(context_for_llm) > 500000: 
    context_for_llm = data_preprocess.normalize_for_overlap(context_for_llm)
    if len(context_for_llm) > 500000:
      if links:
        input_prompt = ["country_name", "modern/ancient/unknown"] 
        if niche_cases: input_prompt += niche_cases 
        reduce_context_for_llm = data_preprocess.build_context_for_llm(texts_reduce, acc, input_prompt, limit_context)
      if reduce_context_for_llm:
        print("reduce context for llm")
        context_for_llm = reduce_context_for_llm
      else:
        print("no reduce context for llm despite>1M")
        context_for_llm = context_for_llm[:limit_context]
  return context_for_llm, linksWithTexts, links

_NEGATION_CUE_PATTERN = re.compile(
    r'\b(?:does not have|does not|did not have|did not|no history of|no evidence of|'
    r'no evidence for|no longer has|absence of|free of|negative for|without|not|no)\s+'
    r'([a-z0-9][a-z0-9 \-]{2,60}?)'
    r'(?=[.,;:)\]]|\band\b|\bor\b|\bbut\b|\bwhile\b|\bwhereas\b|$)',
    re.IGNORECASE,
)

_NEGATION_LEADIN_PATTERN = re.compile(
    r'\b(?:does not have|does not|did not have|did not|no history of|no evidence of|'
    r'no evidence for|no longer has|absence of|free of|negative for|without|not|no)\s*$',
    re.IGNORECASE,
)

_NEGATION_FILLER_WORDS = {
    'a', 'an', 'the', 'any', 'other', 'further', 'additional', 'specific',
    'clear', 'direct', 'obvious', 'known', 'reported', 'documented',
    'available', 'data', 'information', 'details', 'evidence', 'signs',
    'sign', 'this', 'that', 'these', 'those',
}


def _find_negation_contradiction(value: str, explanation: str):
    """Deterministic keyword/negation heuristic (not a semantic re-analysis):
    flags cases where `explanation` clearly negates a condition-phrase
    ("not X", "without X", "no history of X", "does not have X", ...) while
    `value` asserts that same phrase affirmatively elsewhere. Only catches
    direct, literal contradictions -- returns the matched phrase, or None.
    """
    if not value or not explanation:
        return None

    negated_phrases = set()
    for match in _NEGATION_CUE_PATTERN.finditer(explanation):
        phrase = re.sub(r'\s+', ' ', match.group(1).strip(' -'))
        words = [w for w in phrase.split(' ') if w]
        while words and words[0].lower() in _NEGATION_FILLER_WORDS:
            words.pop(0)
        while words and words[-1].lower() in _NEGATION_FILLER_WORDS:
            words.pop()
        phrase = ' '.join(words)
        if len(phrase) >= 4 and len(words) <= 6:
            negated_phrases.add(phrase.lower())

    if not negated_phrases:
        return None

    value_lower = value.lower()
    for phrase in negated_phrases:
        idx = value_lower.find(phrase)
        if idx == -1:
            continue
        preceding = value_lower[max(0, idx - 25):idx]
        if _NEGATION_LEADIN_PATTERN.search(preceding):
            continue
        return phrase

    return None


def _normalize_pass2_json(result: dict, keep_unknown: bool = False) -> dict:
    """Normalize raw Pass-2 JSON to {field: {'value': str, 'explanation': str}}.

    keep_unknown=False (default, single-sample use): drops null/empty/'unknown'
    values, matching the long-standing omit-if-unknown behavior of
    _extract_additional_fields(). keep_unknown=True (batch use): keeps explicit
    'unknown' entries -- required for _extract_additional_fields_batch()'s
    uniform-key-set-per-sample guarantee to survive normalization instead of
    silently collapsing back into per-sample omission.

    Tolerates the model falling back to a flat {field: value} despite
    instructions to use {field: {value, explanation}}.
    """
    cleaned: dict = {}
    skip_vals = {'none', 'null', 'n/a', 'na', 'missing', 'not applicable', ''} | (
        set() if keep_unknown else {'unknown'}
    )
    for k, v in result.items():
        k_str = str(k).strip().lower().replace(' ', '_')
        if not k_str:
            continue
        if isinstance(v, dict):
            v_str = str(v.get('value', '')).strip()
            expl = str(v.get('explanation', '')).strip()
        else:
            v_str = str(v).strip() if v is not None else ''
            expl = ''
        if v_str and v_str.lower() not in skip_vals:
            contradiction_phrase = _find_negation_contradiction(v_str, expl)
            if contradiction_phrase:
                v_str = (
                    f"{v_str} ##SELF-CONTRADICTION: value/explanation disagree "
                    f"(explanation negates '{contradiction_phrase}')"
                )
            cleaned[k_str] = {'value': v_str, 'explanation': expl}
    return cleaned


def _extract_additional_fields(context_text: str, niche_cases: list, standardization_schema: dict = None) -> dict:
    """
    Pass 2 — Generalized metadata extraction across ALL source texts.

    The full multi-source context (NCBI BioSample XML + BioProject + SRA +
    paper text) is passed in so the LLM can:
      a) extract every available attribute, and
      b) detect conflicts when two sources report different values for the
         same field (marked with '##CONFLICT:' in the value string).

    standardization_schema: optional dict {field: {"description", "allowed_values"}}
    -- same shape/source as Pass 1's, falls back to the built-in default schema
    when not supplied, via the same _build_schema_hint() helper Pass 1 uses.

    Returns {field_name: {"value": str, "explanation": str}} — explanation
    contains a one-sentence narrative + a trailing [Sources: ...] tag in the
    same format Pass 1 uses, so downstream code can parse both identically
    instead of leaving Pass-2 fields with no citation at all.
    Safe: always returns a dict (empty on any failure).
    """
    if not context_text or not context_text.strip():
        return {}

    # Fields already handled by Pass 1 — exclude from Pass 2
    exclude_fields = ['country_name', 'modern/ancient/unknown'] + list(niche_cases or [])
    exclude_str = ', '.join(exclude_fields) if exclude_fields else 'none'

    # Keep the full context so nothing is lost; trim only if truly enormous.
    # additional_pipeline.py already caps the combined source text at 800K
    # chars before it reaches here, so this only protects against a single
    # call site bypassing that cap -- it must not re-truncate within that
    # budget, or content past the old 120K cut (e.g. a late-appearing table)
    # is silently dropped from this pass.
    MAX_CHARS = 800000
    context_snippet = context_text if len(context_text) <= MAX_CHARS else context_text[:MAX_CHARS]

    _study_name_hint = (
        "IMPORTANT — for study/dataset-identifier fields: use the paper's own identifying "
        "convention (e.g. first-author surname + publication year, such as 'SmithJ_2021'), or "
        "the actual publication title if no such convention is stated. Do NOT output an "
        "NCBI/ENA accession number (BioProject, BioSample, SRA/ENA study ID, run accession, "
        "etc.) as the answer -- an accession-shaped string is never a valid study/dataset name.\n"
    )

    _disease_hint = (
        "CONTROL DEFINITION: a 'control' sample belongs to the group with NONE of the study's conditions/exposures present (the fully unaffected/reference group) -- "
        "not merely 'not the primary condition being studied.' If you cannot confidently determine full-unaffected status for this sample, output 'unknown' rather than "
        "defaulting to a case/disease label. "
        "Target condition definition: Primary phenotype(s), condition(s), or disease status THAT THIS SPECIFIC SAMPLE ACTUALLY HAS in the study, "
        "as determined by this sample's group assignment or individual clinical status. NOT the study's overall research topic. "
        "CROSS-FIELD CONSISTENCY: several extracted fields may describe the same underlying case/control assignment from different angles -- "
        "whatever you conclude for one such field, every other such field must agree for THIS SAME sample; never report one field as "
        "control/unaffected while another names a condition as if this sample has it, or vice versa."
    )

    _caveat = ("IMPORTANT — per-subject assignment: many studies have multiple participant groups "
        "(e.g. a reference/control group and one or more affected/exposed groups). "
        "Search for a table or supplementary file that maps individual sample identifiers (sample IDs / subject IDs / NCBI BioSample, isolate name, etc. in the NCBI records) "
        "accessions to their specific group. The accession being analysed is shown in 'Prompt N:' above — find its row in that table and extract the group/condition for "
        "THAT SPECIFIC SAMPLE, not the study as a whole. "
        "PRIORITY RULE: a table/section that maps individual sample identifiers to their specific category is stronger evidence than general prose describing the study's groups or conditions as a whole; "
        "when such a table/section exist, cite its matching row, not the prose. "
        "COMPLETENESS: do not stop at the first candidate table/section you find. First check what identifying attribute(s) actually exist on this sample's own record, "
        "whatever they are called (id, subject_id, isolate_name, strain, specimen_code, patient number, or anything else) -- "
        "then check EVERY table/section in the source text that uses a matching or clearly related identifying scheme, "
        "not just the first or most prominent one. Check NCBI BioSample attributes and also paper tables / supplementary metadata tables. "
        "Do NOT report the full list of study groups — report only the group for this individual sample. "
        )

    _schema = standardization_schema if standardization_schema else _get_default_schema()
    schema_hint = _build_schema_hint(niche_cases, _schema)

    generalized_prompt = (
        "You are a scientific metadata extractor specialising in NCBI genomic database records.\n\n"
        "The source text below contains ALL available texts for this sample, separated by "
        "'-----END OF THIS SOURCE <name> ----' markers. Sources may include:\n"
        "  • NCBI BioSample XML (most reliable — look for <Attribute attribute_name='FIELD'>VALUE</Attribute>)\n"
        "  • SRA experiment XML (platform, library strategy, instrument model, etc.)\n"
        "  • BioProject description\n"
        "  • Published paper abstract / full text\n"
        "  • User-uploaded supplementary files\n\n"
        "Your task: extract EVERY metadata attribute that describes the biological sample.\n"
        f"For EACH FIELD: Extract the SAMPLE'S ACTUAL STATUS or INDIVIDUAL VALUE (what applies to THIS sample), "
        f"never the study's general description or objectives. Even for fields defined in terms of 'the study,' prioritize this sample's specific measured/assigned value. "
        f"Only use study-level information as fallback when this sample has no individual value recorded.\n"
        f"{schema_hint}"
        f"{_disease_hint}\n"
        f"{_study_name_hint}"
        f"{_caveat}\n"
        "Scan ALL sources. For EACH field:\n"
        "  - If every source agrees on the same value → output that value.\n"
        "  - If two or more sources report DIFFERENT values → output the most specific value "
        "    AND append '##CONFLICT: source_A=<value_A>, source_B=<value_B>' so conflicts are visible.\n\n"
        "LOOK ESPECIALLY FOR (but extract everything you find):\n"
        "  geo_loc_name, host, tissue, isolation_source, collection_date, sex, age, disease,\n"
        "  treatment, organism, strain, sample_type, body_site, library_strategy, library_source,\n"
        "  library_selection, platform, instrument_model, sequencing_platform, dna_extraction_kit,\n"
        "  lat_lon, env_biome, env_feature, env_material, depth, altitude, temperature, pH,\n"
        "  SRA_accession, BioSample_accession, and any other custom sample attributes.\n\n"
        f"Do NOT include these already-extracted fields: {exclude_str}\n\n"
        "Return ONLY a JSON object mapping each field to an object with TWO keys, explanation and value --\n"
        "write explanation FIRST, decide value only AFTER, based on that reasoning (never decide value before "
        "you've worked out the explanation for that same field):\n"
        '  {"field_name": {"explanation": "<one sentence citing WHERE this came '
        "from, naming the specific source/section/attribute, followed by a "
        "[Sources: <key> (<location>, '<verbatim excerpt <=15 words>')] tag, and for any categorical/group-type "
        "field (disease, condition, status, diagnosis, group, phenotype, health) also ending with an "
        "[ID-match: true|false] tag>\", \"value\": \"<the extracted/best value, decided from the explanation above>\"}}\n"
        "  - Keys  : lowercase field names, underscores for spaces (e.g. 'collection_date')\n"
        "  - explanation : MANDATORY, never blank, written BEFORE value; must include the [Sources: ...] tag using "
        "the exact header from 'The source - <key>:' blocks in the text\n"
        "  - value : the extracted/best value as a non-empty string, consistent with its own explanation above "
        "and with every other field's value for this sample (see CROSS-FIELD CONSISTENCY above)\n"
        "  - Omit fields whose value is null, empty, 'not applicable', 'missing', or 'unknown'\n"
        "  - Preserve the original attribute name from NCBI XML when possible\n\n"
        "Source text:\n"
        "---\n"
        f"{context_snippet}\n"
        "---\n\n"
        "Return ONLY valid JSON. No markdown fences.\n"
        'Example: {"geo_loc_name": {"explanation": '
        "\"BioSample attribute geo_loc_name is 'USA: California'. [Sources: NCBI_biosample (geo_loc_name "
        "attribute, 'USA: California')]\", \"value\": \"USA: California\"}, \"sex\": {\"explanation\": \"Methods section "
        "states male donor. [Sources: https://doi.org/10.1234/x (Methods, 'male donor')]\", \"value\": \"male\"}}"
    )

    try:
        response_text, _ = call_llm_api(generalized_prompt)
        raw = response_text.strip()

        # Strip markdown fences if the model wraps the output despite instructions
        if raw.startswith('```'):
            parts = raw.split('```')
            raw = parts[1] if len(parts) > 1 else raw
            if raw.lower().startswith('json'):
                raw = raw[4:]

        # Find JSON object boundaries in case the model prepends/appends text
        brace_start = raw.find('{')
        brace_end   = raw.rfind('}')
        if brace_start != -1 and brace_end != -1 and brace_end > brace_start:
            raw = raw[brace_start:brace_end + 1]

        result = json.loads(raw.strip())
        return _normalize_pass2_json(result)

    except Exception as e:
        print(f'[_extract_additional_fields] WARNING: generalized extraction failed: {e}')
        return {}


BATCH_PROMPT_CHAR_LIMIT = 800000  # same ceiling _extract_additional_fields() already
                                   # trusts for a single sample's context (its own
                                   # MAX_CHARS above) -- reused here as a simple,
                                   # already-calibrated number rather than inventing
                                   # a new one for batches.


def _split_oversized_batch(prompts: dict, char_limit: int = BATCH_PROMPT_CHAR_LIMIT) -> list:
    """
    Pre-flight size check for a batch about to become one LLM call.

    Hard threshold, not adaptive sizing: estimates size from raw character
    count (no extra API round-trip). If the combined size of `prompts`'
    values fits under `char_limit`, returns [prompts] unchanged. Otherwise
    halves it repeatedly into smaller dicts until each sub-batch fits, or
    it's down to one accession (sent regardless of its own size -- that's
    already the finest granularity the pipeline has).
    """
    total_chars = sum(len(v) for v in prompts.values())
    if total_chars <= char_limit or len(prompts) <= 1:
        return [prompts]
    accs = list(prompts.keys())
    mid = len(accs) // 2
    left  = {a: prompts[a] for a in accs[:mid]}
    right = {a: prompts[a] for a in accs[mid:]}
    return _split_oversized_batch(left, char_limit) + _split_oversized_batch(right, char_limit)


def _extract_additional_fields_batch(contexts: dict, niche_cases: list,
                                      standardization_schema: dict = None) -> dict:
    """
    Pass 2, batched: one LLM call extracts free-form metadata for every
    accession in `contexts` ({acc: context_text}) at once, instead of one
    call per accession.

    Consistency requirement (the reason this isn't just N single-sample
    calls joined together): the model must first determine the UNION of
    attributes present anywhere across the batch's source material, then
    emit that SAME key set for every sample's object -- using an explicit
    "unknown" value for a sample whose own record doesn't state that
    attribute, rather than silently omitting the key the way the
    single-sample _extract_additional_fields() does. Without this, sample A
    could report an attribute (e.g. dna_extraction_kit) that sample B's
    identical shared methods section also states, but B's object just
    omits it -- not because the data differs, but because the model's
    attention landed differently on that generation.

    Returns {acc: {field: {'value': str, 'explanation': str}}} -- 'unknown'
    IS a valid value here (unlike the single-sample function), since it's
    how a uniform key set is expressed for a sample that doesn't have that
    attribute.

    Retries the LLM call + JSON parse once before falling back to an empty
    result for every accession in this batch -- a batch failure costs up to
    BATCH_SIZE accessions' worth of Pass-2 data (vs. 1 for the single-sample
    function), so a cheap retry is worth it here specifically.
    """
    contexts = {a: c for a, c in (contexts or {}).items() if c and c.strip()}
    if not contexts:
        return {}
    accs = list(contexts.keys())
    if len(accs) == 1:
        # No batching benefit or consistency risk at n=1 -- reuse the
        # single-sample function directly instead of a needless JSON
        # restructure.
        return {accs[0]: _extract_additional_fields(contexts[accs[0]], niche_cases, standardization_schema)}

    sub_batches = _split_oversized_batch(contexts)
    if len(sub_batches) > 1:
        print(f"[_extract_additional_fields_batch] {len(contexts)} accession(s) "
              f"({sum(len(v) for v in contexts.values())} chars) exceed the "
              f"{BATCH_PROMPT_CHAR_LIMIT}-char threshold -- splitting into "
              f"{len(sub_batches)} sub-batch(es) of sizes {[len(b) for b in sub_batches]}")
        merged: dict = {}
        for sub in sub_batches:
            merged.update(_extract_additional_fields_batch(sub, niche_cases, standardization_schema))
        return merged

    exclude_fields = ['country_name', 'modern/ancient/unknown'] + list(niche_cases or [])
    exclude_str = ', '.join(exclude_fields) if exclude_fields else 'none'

    _study_name_hint = (
        "IMPORTANT — for study/dataset-identifier fields: use the paper's own identifying "
        "convention (e.g. first-author surname + publication year, such as 'SmithJ_2021'), or "
        "the actual publication title if no such convention is stated. Do NOT output an "
        "NCBI/ENA accession number (BioProject, BioSample, SRA/ENA study ID, run accession, "
        "etc.) as the answer -- an accession-shaped string is never a valid study/dataset name.\n"
    )

    _disease_hint = (
        "CONTROL DEFINITION: a 'control' sample belongs to the group with NONE of the study's conditions/exposures present (the fully unaffected/reference group) -- "
        "not merely 'not the primary condition being studied.' If you cannot confidently determine full-unaffected status for this sample, output 'unknown' rather than "
        "defaulting to a case/disease label. "
        "Target condition definition: Primary phenotype(s), condition(s), or disease status THAT THIS SPECIFIC SAMPLE ACTUALLY HAS in the study, "
        "as determined by this sample's group assignment or individual clinical status. NOT the study's overall research topic. "
        "CROSS-FIELD CONSISTENCY: several extracted fields may describe the same underlying case/control assignment from different angles -- "
        "whatever you conclude for one such field, every other such field must agree for THIS SAME sample; never report one field as "
        "control/unaffected while another names a condition as if this sample has it, or vice versa."
    )

    _caveat = ("IMPORTANT — per-subject assignment: many studies have multiple participant groups "
        "(e.g. a reference/control group and one or more affected/exposed groups). "
        "Search for a table or supplementary file that maps individual sample identifiers (sample IDs / subject IDs / NCBI BioSample, isolate name, etc. in the NCBI records) "
        "accessions to their specific group. Each sample's accession is shown in its own 'SAMPLE N (accession X)' header below -- find its row in that table and extract "
        "the group/condition for THAT SPECIFIC SAMPLE, not the study as a whole. "
        "PRIORITY RULE: a table/section that maps individual sample identifiers to their specific category is stronger evidence than general prose describing the study's groups or conditions as a whole; "
        "when such a table/section exist, cite its matching row, not the prose. "
        "COMPLETENESS: do not stop at the first candidate table/section you find. First check what identifying attribute(s) actually exist on each sample's own record, "
        "whatever they are called (id, subject_id, isolate_name, strain, specimen_code, patient number, or anything else) -- "
        "then check EVERY table/section in the source text that uses a matching or clearly related identifying scheme, "
        "not just the first or most prominent one. Check NCBI BioSample attributes and also paper tables / supplementary metadata tables. "
        "Do NOT report the full list of study groups — report only the group for each individual sample. "
        )

    _schema = standardization_schema if standardization_schema else _get_default_schema()
    schema_hint = _build_schema_hint(niche_cases, _schema)

    per_sample_blocks = [
        f"--- SAMPLE {i+1} (accession {acc.split('.')[0] if acc else acc}) SOURCE TEXT ---\n{contexts[acc]}\n"
        for i, acc in enumerate(accs)
    ]
    samples_text = "\n".join(per_sample_blocks)

    generalized_prompt = (
        "You are a scientific metadata extractor specialising in NCBI genomic database records.\n\n"
        f"Below are source texts for {len(accs)} DIFFERENT biological samples from the same study/paper, "
        "each in its own 'SAMPLE N (accession X)' block. Within each sample's block, its texts are "
        "separated by '-----END OF THIS SOURCE <name> ----' markers. Sources may include:\n"
        "  • NCBI BioSample XML (most reliable — look for <Attribute attribute_name='FIELD'>VALUE</Attribute>)\n"
        "  • SRA experiment XML (platform, library strategy, instrument model, etc.)\n"
        "  • BioProject description\n"
        "  • Published paper abstract / full text\n"
        "  • User-uploaded supplementary files\n\n"
        "Your task: extract EVERY metadata attribute that describes EACH sample.\n"
        f"For EACH FIELD, for EACH sample: Extract THAT SAMPLE'S ACTUAL STATUS or INDIVIDUAL VALUE (what "
        f"applies to THAT specific sample), never the study's general description or objectives. Only use "
        f"study-level information as fallback when a sample has no individual value recorded.\n"
        f"{schema_hint}"
        f"{_disease_hint}\n"
        f"{_study_name_hint}"
        f"{_caveat}\n"
        "CONSISTENCY REQUIREMENT — do this in two steps:\n"
        "  STEP 1: scan ALL samples' source text and determine the UNION of every metadata attribute "
        "present anywhere in this batch.\n"
        "  STEP 2: for EVERY sample, output that SAME set of attribute keys -- if a specific sample's own "
        "source text does not state a given attribute (even though another sample in this batch does), "
        "output \"value\": \"unknown\" for that key on that sample rather than omitting the key. Do not let "
        "one sample's object have a different key set than another's.\n"
        "Scan ALL sources for each sample. For EACH field on EACH sample:\n"
        "  - If every source (for that sample) agrees on the same value → output that value.\n"
        "  - If two or more sources report DIFFERENT values for that sample → output the most specific value "
        "    AND append '##CONFLICT: source_A=<value_A>, source_B=<value_B>' so conflicts are visible.\n\n"
        "LOOK ESPECIALLY FOR (but extract everything you find):\n"
        "  geo_loc_name, host, tissue, isolation_source, collection_date, sex, age, disease,\n"
        "  treatment, organism, strain, sample_type, body_site, library_strategy, library_source,\n"
        "  library_selection, platform, instrument_model, sequencing_platform, dna_extraction_kit,\n"
        "  lat_lon, env_biome, env_feature, env_material, depth, altitude, temperature, pH,\n"
        "  SRA_accession, BioSample_accession, and any other custom sample attributes.\n\n"
        f"Do NOT include these already-extracted fields: {exclude_str}\n\n"
        "Return ONLY a JSON object with one top-level key per sample, using the EXACT accession string "
        "shown in that sample's 'SAMPLE N (accession X)' header as the key, mapping to an object with one "
        "key per attribute, each an object with TWO keys, explanation and value -- write explanation FIRST, "
        "decide value only AFTER, based on that reasoning:\n"
        '  {"<accession>": {"field_name": {"explanation": "<one sentence citing WHERE this came '
        "from, naming the specific source/section/attribute, followed by a "
        "[Sources: <key> (<location>, '<verbatim excerpt <=15 words>')] tag, and for any categorical/group-type "
        "field (disease, condition, status, diagnosis, group, phenotype, health) also ending with an "
        "[ID-match: true|false] tag; if value is \"unknown\" per the CONSISTENCY REQUIREMENT, explain why "
        "(e.g. 'not stated in this sample's own record')>\", \"value\": \"<the extracted/best value, or "
        "\\\"unknown\\\" per the CONSISTENCY REQUIREMENT>\"}, ...}, ...}\n"
        "  - Keys  : lowercase field names, underscores for spaces (e.g. 'collection_date')\n"
        "  - explanation : MANDATORY even when value is \"unknown\", never blank, written BEFORE value; "
        "must include the [Sources: ...] tag using the exact header from 'The source - <key>:' blocks "
        "in the text when value is not \"unknown\"\n"
        "  - value : the extracted/best value as a non-empty string, or \"unknown\" per the CONSISTENCY "
        "REQUIREMENT -- never omit a key the union in STEP 1 identified\n"
        "  - Preserve the original attribute name from NCBI XML when possible\n\n"
        "Source text:\n"
        "---\n"
        f"{samples_text}\n"
        "---\n\n"
        "Return ONLY valid JSON. No markdown fences.\n"
        'Example for 2 samples: {"SAMN001": {"geo_loc_name": {"explanation": '
        "\"BioSample attribute geo_loc_name is 'USA: California'. [Sources: NCBI_biosample (geo_loc_name "
        "attribute, 'USA: California')]\", \"value\": \"USA: California\"}}, \"SAMN002\": {\"geo_loc_name\": "
        "{\"explanation\": \"Not stated in this sample's own BioSample record or the shared paper text.\", "
        "\"value\": \"unknown\"}}}"
    )

    result = None
    for _attempt in range(2):
        try:
            response_text, _ = call_llm_api(generalized_prompt)
            raw = response_text.strip()
            if raw.startswith('```'):
                parts = raw.split('```')
                raw = parts[1] if len(parts) > 1 else raw
                if raw.lower().startswith('json'):
                    raw = raw[4:]
            brace_start = raw.find('{')
            brace_end = raw.rfind('}')
            if brace_start != -1 and brace_end != -1 and brace_end > brace_start:
                raw = raw[brace_start:brace_end + 1]
            result = json.loads(raw.strip())
            break
        except Exception as e:
            if _attempt == 0:
                print(f'[_extract_additional_fields_batch] attempt 1 failed ({e}) -- retrying once')
                continue
            # Both identical attempts failed. A repeat failure at
            # (near-)identical output position across both attempts (as seen
            # live on PRJNA976261: char 24772 then char 24717, both landing
            # just past the same JSON delimiter) indicates the OUTPUT this
            # batch demands -- samples x union-of-attributes, inflated
            # further by the uniform-key-set requirement -- is genuinely too
            # large for this call's fixed output budget, not transient noise
            # a third identical retry would fix. Halve the batch and recurse
            # instead, reusing the same halving strategy
            # _split_oversized_batch already uses for oversized input.
            if len(accs) > 1:
                mid = len(accs) // 2
                left_accs, right_accs = accs[:mid], accs[mid:]
                print(f'[_extract_additional_fields_batch] both attempts failed ({e}) -- '
                      f'output likely too large for {len(accs)} accession(s); splitting into '
                      f'{len(left_accs)} + {len(right_accs)} and retrying each half')
                left  = _extract_additional_fields_batch(
                    {a: contexts[a] for a in left_accs}, niche_cases, standardization_schema)
                right = _extract_additional_fields_batch(
                    {a: contexts[a] for a in right_accs}, niche_cases, standardization_schema)
                return {**left, **right}
            print(f'[_extract_additional_fields_batch] WARNING: both attempts failed for the '
                  f'single remaining accession {accs[0]!r}, last error: {e}')
            return {acc: {} for acc in accs}

    # Tolerant accession matching (model may not echo the key byte-identical --
    # whitespace/case/a dropped version suffix), same spirit as
    # split_batched_llm_response's number-primary/position-fallback approach.
    out: dict = {}
    for acc in accs:
        acc_cleaned = acc.split('.')[0] if acc else acc
        match_key = next(
            (rk for rk in result if rk == acc or rk.split('.')[0] == acc_cleaned), None
        )
        if match_key is None:
            print(f"[_extract_additional_fields_batch] no result object for {acc!r} in batch "
                  f"response (model returned keys: {list(result.keys())}) -- treating as empty")
            out[acc] = {}
            continue
        raw_fields = result[match_key] if isinstance(result[match_key], dict) else {}
        out[acc] = _normalize_pass2_json(raw_fields, keep_unknown=True)
    return out


async def query_document_info(niche_cases, saveLinkFolder, llm_api_function, prompts,
                              standardization_schema=None):
    """
    Queries the document using a hybrid approach:
    1. Local structured lookup (fast, cheap, accurate for known patterns).
    2. RAG with semantic search and LLM (general, flexible, cost-optimized).
    """
    print("inside the model.query_doc_info")

    # ── Pre-flight size check ───────────────────────────────────────────────
    # Hard threshold, not adaptive sizing: batching doesn't currently dedupe
    # shared paper text across accessions in the same batch (deferred
    # follow-up), so a batch of up to BATCH_SIZE accessions could still be
    # oversized for a paper with a large full-text/supplementary load. Split
    # and recurse rather than sending an oversized request and letting it
    # fail/degrade unpredictably.
    sub_batches = _split_oversized_batch(prompts)
    if len(sub_batches) > 1:
        print(f"[query_document_info] batch of {len(prompts)} accession(s) "
              f"({sum(len(v) for v in prompts.values())} chars) exceeds "
              f"{BATCH_PROMPT_CHAR_LIMIT}-char safety threshold -- splitting into "
              f"{len(sub_batches)} sub-batch(es) of sizes {[len(b) for b in sub_batches]}")
        merged_outputs = {}
        for sub in sub_batches:
            merged_outputs.update(await query_document_info(
                niche_cases, saveLinkFolder, llm_api_function, sub,
                standardization_schema=standardization_schema))
        return merged_outputs

    outputs, links, accession_found_in_text = {}, [], False

    genai.configure(api_key=os.getenv("NEW_GOOGLE_API_KEY") or os.getenv("GOOGLE_API_KEY") or os.getenv("NEW_GEMINI_API"))
    # Gemini 2.5 Flash-Lite pricing per 1,000 tokens
    PRICE_PER_1K_INPUT_LLM = 0.00010      # $0.10 per 1M input tokens
    PRICE_PER_1K_OUTPUT_LLM = 0.00040     # $0.40 per 1M output tokens
    
    # Embedding-001 pricing per 1,000 input tokens
    PRICE_PER_1K_EMBEDDING_INPUT = 0.00015  # $0.15 per 1M input tokens
    global_llm_model_for_counting_tokens = genai.GenerativeModel("gemini-2.5-flash-lite")#('gemini-1.5-flash-latest')

    # Determine fields to ask LLM for and output format based on what's known/needed
    method_used = 'rag_llm' # Will be updated based on the method that yields a result
    # Only add country_name/sample_type defaults when niche_cases doesn't already cover them.
    # This avoids duplicate fields and fixes parsing when niche_cases includes "country".
    _nc_lower = [nc.lower() for nc in (niche_cases or [])]
    _has_country_field = any('country' in nc for nc in _nc_lower)
    _has_type_field = any(nc in ('sample_type', 'modern/ancient/unknown', 'modern', 'ancient',
                                  'sample type', 'sample_source_type') for nc in _nc_lower)
    _default_fields = []
    if not _has_country_field:
        _default_fields.append("country_name")
    if not _has_type_field:
        _default_fields.append("modern/ancient/unknown")
    if niche_cases:
        _all_fields = _default_fields + list(niche_cases)
    elif _default_fields:
        _all_fields = _default_fields
    else:
        _all_fields = ["country_name", "modern/ancient/unknown"]
    output_format_str = ", ".join(_all_fields)
    # Calculate embedding cost for the primary query word
    total_query_cost, current_embedding_cost = 0, 0
    created_prompts = multi_prompts(prompts, output_format_str, niche_cases=niche_cases,
                                    prompt_template="default",
                                    standardization_schema=standardization_schema)
    print("done create prompt and length: ", len(created_prompts))
    prompt_for_llm = []
    for acc in created_prompts:
      outputs[acc] = {"predicted_output":"",
                      "method_used": method_used,
                      "total_query_cost":None,
                      "links": [],
                      "accession_found_in_text":created_prompts[acc][1],
                      }
      prompt_for_llm.append(created_prompts[acc][0])  
    
    prompt_for_llm = "\n".join(prompt_for_llm) #there is only 1 prompt created #+ "\n" + "Give answer for each prompt"
    print("length of prompt: ", len(prompt_for_llm))
    print("use 2.5 flash gemini")
    llm_response_text, model_instance = call_llm_api(prompt_for_llm)
    print("\n--- DEBUG INFO FOR RAG ---")
    print("Retrieved Context Sent to LLM (first 500 chars):")
    print(prompt_for_llm[:500] + "..." if len(prompt_for_llm) > 500 else prompt_for_llm)
    print("\nRaw LLM Response:")
    print(llm_response_text)
    print("--- END DEBUG INFO ---")
        
    llm_cost = 0
    if model_instance:
        try:
            input_llm_tokens = global_llm_model_for_counting_tokens.count_tokens(prompt_for_llm).total_tokens
            output_llm_tokens = global_llm_model_for_counting_tokens.count_tokens(llm_response_text).total_tokens
            print(f"  DEBUG: LLM Input tokens: {input_llm_tokens}")
            print(f"  DEBUG: LLM Output tokens: {output_llm_tokens}")
            llm_cost = (input_llm_tokens / 1000) * PRICE_PER_1K_INPUT_LLM + \
                       (output_llm_tokens / 1000) * PRICE_PER_1K_OUTPUT_LLM
            print(f"  DEBUG: Estimated LLM cost: ${llm_cost:.6f}")
        except Exception as e:
            print(f"  DEBUG: Error counting LLM tokens: {e}")
            llm_cost = 0

    total_query_cost += current_embedding_cost + llm_cost
    print(f"  DEBUG: Total estimated cost for this RAG query: ${total_query_cost:.6f}")
    
    list_accs = list(prompts.keys())
    segments = split_batched_llm_response(llm_response_text, list_accs)
    for acc in list_accs:
      metadata_list = parse_multi_sample_llm_output(segments.get(acc, ""), output_format_str)
      print(metadata_list)
      again_output_format, general_knowledge_prompt = "", ""
      output_acc = {}
      # NOTE: an unknown-field retry used to re-run getMoreInfoForAcc() here,
      # but both pipeline.py and additional_pipeline.py already call it
      # (smart web search) before building prompts[acc] -- that search result
      # is already folded into this context, so re-running it found nothing
      # new and only doubled the search cost/time. Use the context as-is.
      context_for_llm = prompts[acc]
      # Collect unknown fields and known fields in one pass
      unknown_fields = []
      for key in metadata_list:
        answer = metadata_list[key]["answer"]
        if answer.lower() in ("unknown", "unspecified", "could not get response from llm api.", "undefined"):
          unknown_fields.append(key)
        else:
          output_acc[key] = metadata_list[key]

      # Unknown fields after the first LLM pass stay as "unknown" — no retry.
      # Re-running the same niche fields against the same context rarely recovers new info
      # and costs an extra LLM call per sample.
      if unknown_fields:
        print(f"{len(unknown_fields)} field(s) returned unknown — keeping as-is: {unknown_fields}")
        for uf in unknown_fields:
          output_acc[uf] = {"answer": "unknown", f"{uf}_explanation": "unknown"}
      # ── LLM-based standardization pass ───────────────────────────────────
      # Run after extraction; maps free-text values to canonical schema values.
      if standardization_schema and output_acc:
          try:
              extracted_flat = {
                  k: output_acc[k]["answer"]
                  for k in output_acc
                  if isinstance(output_acc[k], dict) and output_acc[k].get("answer", "").lower() not in ("", "unknown")
              }
              if extracted_flat:
                  standardized = standardize_with_llm(extracted_flat, standardization_schema, acc)
                  for field, std_val in standardized.items():
                      if field in output_acc and std_val:
                          output_acc[field]["answer"] = std_val
                  print(f"[Standardization] {acc}: {standardized}")
          except Exception as _std_err:
              print(f"[Standardization] WARNING: {_std_err}")

      outputs[acc]["predicted_output"] = output_acc
      outputs[acc]["total_query_cost"] = total_query_cost

      print("total cost: ", total_query_cost)

    # ── PASS 2: generalized extraction of ALL additional metadata ───────────
    # Batched across every accession in this call (one LLM call for the whole
    # batch, not one per accession) -- moved out of the per-acc loop above.
    # Uses all source text to extract every metadata attribute mentioned.
    try:
        predefined_keys = set(['country_name', 'modern/ancient/unknown']
                               + list(niche_cases or []))
        # Always use the original multi-source text (prompts[acc]) so that
        # BioSample XML attributes and paper text are both available.
        pass2_contexts = {acc: prompts.get(acc, "") for acc in list_accs}
        all_additional_by_acc = _extract_additional_fields_batch(
            pass2_contexts, niche_cases or [], standardization_schema=standardization_schema)
        for acc in list_accs:
            additional_only = {
                k: v for k, v in (all_additional_by_acc.get(acc) or {}).items()
                if k not in predefined_keys
            }
            outputs[acc]['_additional_fields'] = additional_only
            print(f'[Pass 2] {acc}: {len(additional_only)} additional fields -> '
                  f'{list(additional_only.keys())}')
    except Exception as _pass2_err:
        print(f'[Pass 2] WARNING: batch failed: {_pass2_err}')
        for acc in list_accs:
            outputs[acc]['_additional_fields'] = {}
    # ── END PASS 2 ────────────────────────────────────────────────────────

    for acc in list_accs:
      print(f"total output of {acc}: {outputs[acc]}")
    return outputs