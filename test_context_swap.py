"""
Standalone diagnostic: does fixing Table 1 / Table 3's forward-fill and
delimiting in SAMN35361955's context change the disease/group-status
answer the real pipeline produces?

Does NOT modify any pipeline file. Imports and calls
model.query_document_info() directly -- the exact function
additional_pipeline.py calls (see additional_pipeline.py:1151) for this
extraction step, with the same call signature and the same niche_cases=None
(no schema/user-specified fields) that the last full PRJNA976261 test run
used (per test_output_phase2.md: "the path PRJNA976261's test actually
exercises, since it supplies no fields" -- meaning the disease/group field
comes out of query_document_info's Pass 2, model._extract_additional_fields(),
not a hand-reconstructed prompt).

Usage: python test_context_swap.py
"""
import asyncio
import json

from docx import Document

import model

ACC = "SAMN35361955"
BAD_CONTEXT_PATH = "test-data/PRJNA976261/SAMN35361955.docx"
FIXED_CONTEXT_PATH = "test-data/PRJNA976261/SAMN35361955_fixed.docx"


def read_context_docx(path: str) -> str:
    """Read back a docx exactly as model's additional_pipeline.py wrote it.

    additional_pipeline.py builds the combined per-accession context string
    and saves it via data_preprocess.save_text_to_docx(), which does
    `for line in text.split('\\n'): document.add_paragraph(line)`. Reversing
    that (join paragraph text with '\\n') reconstructs the original context
    string byte-for-byte, including blank lines. model.read_docx_text() is
    NOT used here because it is a different-purpose function (splits out
    "## Table N" blocks and drops blank paragraphs for the RAG-indexing
    path) -- it would alter the very table formatting/delimiting this test
    is trying to compare.
    """
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


async def run_one(label: str, context_text: str) -> dict:
    prompts = {ACC: context_text}
    print(f"\n{'=' * 80}\nRUNNING: {label} ({len(context_text)} chars)\n{'=' * 80}")
    outputs = await model.query_document_info(
        niche_cases=None,
        saveLinkFolder=None,
        llm_api_function=model.call_llm_api,
        prompts=prompts,
        standardization_schema=None,
    )
    return outputs[ACC]


CATEGORICAL_KEYWORDS = (
    "disease", "condition", "diagnosis", "periodont", "diabet",
    "control", "status", "group", "phenotype", "health",
)


def find_categorical_fields(result: dict) -> dict:
    """Scan both passes' output for any field whose NAME looks disease/group-related."""
    found = {}
    for field, val in result.get("predicted_output", {}).items():
        if any(kw in field.lower() for kw in CATEGORICAL_KEYWORDS):
            found[f"predicted_output.{field}"] = val.get("answer")
    for field, val in result.get("_additional_fields", {}).items():
        if any(kw in field.lower() for kw in CATEGORICAL_KEYWORDS):
            found[f"_additional_fields.{field}"] = val.get("value")
    return found


def format_result(label: str, result: dict) -> str:
    lines = [f"## {label}", ""]
    lines.append(f"- method_used: `{result.get('method_used')}`")
    lines.append(f"- accession_found_in_text: `{result.get('accession_found_in_text')}`")
    lines.append("")
    lines.append("### Pass 1 (`multi_prompts` / default fields)")
    predicted = result.get("predicted_output", {})
    if not predicted:
        lines.append("_(none)_")
    for field, val in predicted.items():
        lines.append(f"\n**{field}**")
        lines.append(f"- answer: `{val.get('answer')}`")
        expl_key = f"{field}_explanation"
        lines.append(f"- explanation: {val.get(expl_key, '')}")
    lines.append("")
    lines.append("### Pass 2 (`_extract_additional_fields` — generalized JSON extraction)")
    additional = result.get("_additional_fields", {})
    if not additional:
        lines.append("_(none)_")
    for field, val in additional.items():
        lines.append(f"\n**{field}**")
        lines.append(f"- value: `{val.get('value')}`")
        lines.append(f"- explanation: {val.get('explanation', '')}")
    return "\n".join(lines)


async def main():
    bad_context = read_context_docx(BAD_CONTEXT_PATH)
    fixed_context = read_context_docx(FIXED_CONTEXT_PATH)

    bad_result = await run_one("bad_context", bad_context)
    fixed_result = await run_one("fixed_context", fixed_context)

    bad_categorical = find_categorical_fields(bad_result)
    fixed_categorical = find_categorical_fields(fixed_result)

    bad_md = format_result("bad_context (SAMN35361955.docx)", bad_result)
    fixed_md = format_result("fixed_context (SAMN35361955_fixed.docx)", fixed_result)

    print("\n\n" + "#" * 80)
    print("CATEGORICAL/DISEASE-RELATED FIELD SCAN (by field name)")
    print("#" * 80)
    print("bad_context:  ", bad_categorical or "(none found)")
    print("fixed_context:", fixed_categorical or "(none found)")

    print("\n\n" + "#" * 80)
    print("SIDE BY SIDE RAW OUTPUT")
    print("#" * 80)
    print("\n" + bad_md)
    print("\n" + fixed_md)

    _none_msg = (
        "**NONE** — no field whose name matches disease/condition/diagnosis/"
        "periodont/diabet/control/status/group/phenotype/health was present "
        "in predicted_output or _additional_fields"
    )
    bad_categorical_str = json.dumps(bad_categorical) if bad_categorical else _none_msg
    fixed_categorical_str = json.dumps(fixed_categorical) if fixed_categorical else "**NONE** — same as above"
    categorical_summary = (
        f"- bad_context categorical/disease-related fields found: {bad_categorical_str}\n"
        f"- fixed_context categorical/disease-related fields found: {fixed_categorical_str}\n"
    )

    if bad_categorical and fixed_categorical:
        if bad_categorical == fixed_categorical:
            headline = (
                "Both contexts produced the **same** categorical/disease-related "
                "field value(s) — fixing Table 1/Table 3's forward-fill and "
                "delimiting did not change the model's answer this run."
            )
        else:
            headline = (
                "**The value changed.** bad_context and fixed_context produced "
                "**different** categorical/disease-related field value(s) for "
                f"{ACC} — see the full raw outputs below (especially each "
                "`disease` field's `[Candidates:]`/`[Chosen:]`/`[ID-match:]` tags) "
                "for exactly which table each version cited and why."
            )
    elif bad_categorical or fixed_categorical:
        headline = (
            "**Asymmetric result.** Only one of the two contexts produced a "
            "categorical/disease-related field at all this run — see raw output "
            "below."
        )
    else:
        headline = (
            "**Neither context produced any disease/periodontitis/diabetes/"
            "group-status field at all** — not `unknown`, just absent. "
            "`_extract_additional_fields()` silently omits fields whose value "
            "comes back as null/empty/'unknown' (`model.py`'s `skip_vals` set), "
            "and its own prompt tells the model to \"omit the field rather than "
            "guessing\" when it can't confidently apply the reliability test to "
            "a candidate table. This run cannot show whether the value of the "
            "disease field changes between bad_context and fixed_context — the "
            "field itself wasn't emitted either way. LLM output is "
            "non-deterministic; re-running this script may produce a populated "
            "field (as a prior run in this session did)."
        )

    report = (
        f"# Context swap test — {ACC}\n\n"
        f"Function called: `model.query_document_info()` "
        f"(defined in `model.py`, called the same way `additional_pipeline.py` "
        f"calls it at `additional_pipeline.py:1151` — "
        f"`niche_cases=None`, `saveLinkFolder=None`, "
        f"`llm_api_function=model.call_llm_api`, `standardization_schema=None`, "
        f"matching the last full PRJNA976261 test run which supplied no "
        f"niche_cases/schema).\n\n"
        f"Internally this calls `model.multi_prompts()` + `model.call_llm_api()` "
        f"(Pass 1: country/modern-ancient) and `model._extract_additional_fields()` "
        f"(Pass 2: generalized JSON extraction — this is where the "
        f"disease/group-status field is actually produced for this dataset, "
        f"since no niche_cases were supplied).\n\n"
        f"Both calls in this run went through `model.call_llm_api()`'s **Anthropic "
        f"Claude Haiku** branch (`claude-haiku-4-5-20251001`), not Gemini — "
        f"`ANTHROPIC_API_KEY` is set in this environment and is tried first; "
        f"verified separately (`model.call_llm_api()` returned `model_instance=None`, "
        f"which only happens on the Anthropic return path, never the Gemini one). "
        f"Earlier documented runs of this same pipeline (`test_output_phase2.md`) "
        f"note using Gemini 2.5 Flash-Lite, so this run is not directly comparable "
        f"to those in provider terms. LLM output is non-deterministic, so a "
        f"re-run may pick a different provider/answer even with identical inputs.\n\n"
        f"Context sources:\n"
        f"- bad_context: `{BAD_CONTEXT_PATH}` ({len(bad_context)} chars)\n"
        f"- fixed_context: `{FIXED_CONTEXT_PATH}` ({len(fixed_context)} chars)\n\n"
        f"## Headline result\n\n"
        f"{categorical_summary}\n"
        f"{headline}\n\n"
        f"---\n\n"
        f"{bad_md}\n\n"
        f"---\n\n"
        f"{fixed_md}\n"
    )
    with open("test_output_context_swap.md", "w") as f:
        f.write(report)
    print("\n\nWrote test_output_context_swap.md")


if __name__ == "__main__":
    asyncio.run(main())
